#!/usr/bin/env python3
"""
Comprehensive Resumé Analyzer - Deep examination of all Resumés in Network folder
Analyzes structure, formatting, tables, colors, fonts, tone of voice, and style
"""

import os
import sys
from datetime import datetime
from pathlib import Path
import re
from collections import defaultdict, Counter
import json

def analyze_docx_structure(docx_path):
    """Deep analysis of DOCX structure, formatting, tables, and style"""
    try:
        from docx import Document
        from docx.shared import RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document(docx_path)
        
        analysis = {
            'filename': os.path.basename(docx_path),
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'sections': len(doc.sections),
            'fonts_used': [],
            'font_sizes': [],
            'colors_used': [],
            'bold_usage': 0,
            'italic_usage': 0,
            'table_structures': [],
            'text_content': [],
            'header_styles': [],
            'alignment_patterns': [],
            'spacing_info': {},
            'first_page_content': [],
            'boilerplate_text': [],
        }
        
        # Analyze sections and page setup
        if doc.sections:
            section = doc.sections[0]
            analysis['page_width'] = section.page_width.inches if section.page_width else None
            analysis['page_height'] = section.page_height.inches if section.page_height else None
            analysis['margin_top'] = section.top_margin.inches if section.top_margin else None
            analysis['margin_bottom'] = section.bottom_margin.inches if section.bottom_margin else None
            analysis['margin_left'] = section.left_margin.inches if section.left_margin else None
            analysis['margin_right'] = section.right_margin.inches if section.right_margin else None
        
        # Analyze paragraphs
        for i, para in enumerate(doc.paragraphs[:50]):  # First 50 paragraphs
            if para.text.strip():
                para_info = {
                    'text': para.text.strip()[:200],  # First 200 chars
                    'style': para.style.name if para.style else 'None',
                    'alignment': str(para.alignment) if para.alignment else 'None',
                }
                
                # Analyze runs for formatting
                for run in para.runs:
                    if run.font.name:
                        analysis['fonts_used'].append(run.font.name)
                    if run.font.size:
                        analysis['font_sizes'].append(run.font.size.pt)
                    if run.font.color and run.font.color.rgb:
                        analysis['colors_used'].append(str(run.font.color.rgb))
                    if run.bold:
                        analysis['bold_usage'] += 1
                    if run.italic:
                        analysis['italic_usage'] += 1
                
                if i < 10:  # First 10 paragraphs
                    analysis['first_page_content'].append(para_info)
                
                analysis['text_content'].append(para.text.strip())
        
        # Analyze tables
        for table_idx, table in enumerate(doc.tables):
            table_info = {
                'index': table_idx,
                'rows': len(table.rows),
                'columns': len(table.columns),
                'cells': [],
                'structure_pattern': [],
            }
            
            # Analyze first few rows
            for row_idx, row in enumerate(table.rows[:10]):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text[:100] if cell_text else '')
                table_info['structure_pattern'].append(row_data)
            
            analysis['table_structures'].append(table_info)
        
        return analysis, True
        
    except Exception as e:
        return {'error': str(e), 'filename': os.path.basename(docx_path)}, False

def analyze_resume_naming_pattern(filename):
    """Analyze naming pattern of Resumé file"""
    patterns = {
        'has_accent': 'Resumé' in filename,
        'has_synergie': 'Synergie' in filename or 'synergie' in filename,
        'has_projectmanagement': 'projectmanagement' in filename,
        'has_date': bool(re.search(r'\d{4}', filename)),
        'has_underscore': '_' in filename,
        'extension': os.path.splitext(filename)[1].lower()
    }
    
    # Try to extract name from filename
    name_match = re.search(r'projectmanagement[_\s]+(.+?)\.(pdf|docx)', filename, re.IGNORECASE)
    if name_match:
        patterns['extracted_name'] = name_match.group(1)
    
    return patterns

def detect_tone_of_voice(text_content):
    """Analyze tone of voice and writing style"""
    if not text_content:
        return {}
    
    full_text = ' '.join(text_content).lower()
    
    tone_analysis = {
        'formal_indicators': 0,
        'technical_terms': 0,
        'action_verbs': 0,
        'first_person': 0,
        'third_person': 0,
        'passive_voice': 0,
        'bullet_points': 0,
    }
    
    # Formal indicators
    formal_words = ['desbetreffend', 'betreffende', 'zorgdragen', 'ten behoeve van', 
                    'in samenwerking met', 'verantwoordelijk voor', 'ervaring met']
    for word in formal_words:
        tone_analysis['formal_indicators'] += full_text.count(word)
    
    # Technical terms
    tech_terms = ['projectmanagement', 'planning', 'risicomanagement', 'stakeholder',
                  'scope', 'budget', 'kwaliteit', 'mijlpaal', 'deliverable']
    for term in tech_terms:
        tone_analysis['technical_terms'] += full_text.count(term)
    
    # Action verbs
    action_verbs = ['begeleid', 'georganiseerd', 'gerealiseerd', 'geïmplementeerd',
                    'ontwikkeld', 'gecoördineerd', 'geleid', 'beheerd']
    for verb in action_verbs:
        tone_analysis['action_verbs'] += full_text.count(verb)
    
    # First vs third person
    tone_analysis['first_person'] = full_text.count(' ik ') + full_text.count(' mij ')
    tone_analysis['third_person'] = full_text.count(' hij ') + full_text.count(' zij ')
    
    # Bullet points
    for line in text_content:
        if line.strip().startswith('•') or line.strip().startswith('-'):
            tone_analysis['bullet_points'] += 1
    
    return tone_analysis

def identify_standard_sections(text_content):
    """Identify standard sections in Resumé"""
    sections_found = []
    
    section_patterns = {
        'Werkervaring': [r'werkervaring', r'work\s+experience'],
        'Opleiding': [r'opleiding', r'education', r'opleidingen'],
        'Cursussen': [r'cursussen', r'courses', r'training'],
        'Projecten': [r'projecten', r'projects'],
        'Vaardigheden': [r'vaardigheden', r'skills'],
        'Competenties': [r'competenties', r'competencies'],
        'Certificaten': [r'certificaten', r'certificates'],
        'Talen': [r'talen', r'languages'],
    }
    
    for i, text in enumerate(text_content):
        if len(text) < 50 and text.strip():  # Likely a header
            for section_name, patterns in section_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, text.lower()):
                        sections_found.append({
                            'section': section_name,
                            'line': i,
                            'text': text.strip()
                        })
                        break
    
    return sections_found

def analyze_header_footer(doc):
    """Analyze header and footer content"""
    try:
        from docx import Document
        
        header_footer_info = {
            'has_header': False,
            'has_footer': False,
            'header_content': [],
            'footer_content': [],
        }
        
        # Check sections for headers/footers
        for section in doc.sections:
            # Header
            header = section.header
            if header:
                header_text = []
                for para in header.paragraphs:
                    if para.text.strip():
                        header_text.append(para.text.strip())
                if header_text:
                    header_footer_info['has_header'] = True
                    header_footer_info['header_content'].extend(header_text)
            
            # Footer
            footer = section.footer
            if footer:
                footer_text = []
                for para in footer.paragraphs:
                    if para.text.strip():
                        footer_text.append(para.text.strip())
                if footer_text:
                    header_footer_info['has_footer'] = True
                    header_footer_info['footer_content'].extend(footer_text)
        
        return header_footer_info
    except:
        return {'has_header': False, 'has_footer': False}

def scan_all_resumes(base_path):
    """Scan all Resumé files in the Network folder"""
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Starting comprehensive Resumé scan...")
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Base path: {base_path}")
    
    # Install dependencies if needed
    try:
        from docx import Document
    except ImportError:
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Installing python-docx...")
        os.system("pip install python-docx -q")
    
    resume_pattern = re.compile(r'Resum[eé]_Synergie', re.IGNORECASE)
    all_resume_analyses = []
    
    stats = {
        'total_resumes_found': 0,
        'total_resumes_analyzed': 0,
        'docx_count': 0,
        'pdf_count': 0,
        'fonts_frequency': defaultdict(int),
        'font_sizes_frequency': defaultdict(int),
        'colors_frequency': defaultdict(int),
        'table_counts': defaultdict(int),
        'common_structures': [],
        'naming_patterns': defaultdict(int),
    }
    
    # Walk through all folders
    for root, dirs, files in os.walk(base_path):
        # Skip template folder
        if 'STANDAARD RESUME' in root:
            continue
        
        person_name = os.path.basename(root)
        
        for file in files:
            # Check if it's a Resumé file
            if not resume_pattern.search(file):
                continue
            
            stats['total_resumes_found'] += 1
            full_path = os.path.join(root, file)
            
            # Only analyze DOCX files for detailed structure
            if file.lower().endswith('.docx'):
                stats['docx_count'] += 1
                
                print(f"[{datetime.now().strftime('%H:%M:%S')}] Analyzing: {person_name} - {file}")
                
                # Analyze structure
                analysis, success = analyze_docx_structure(full_path)
                
                if success:
                    analysis['person'] = person_name
                    analysis['filepath'] = full_path
                    
                    # Analyze naming
                    naming = analyze_resume_naming_pattern(file)
                    analysis['naming_pattern'] = naming
                    
                    # Analyze tone
                    if analysis.get('text_content'):
                        tone = detect_tone_of_voice(analysis['text_content'])
                        analysis['tone_of_voice'] = tone
                        
                        # Identify sections
                        sections = identify_standard_sections(analysis['text_content'])
                        analysis['sections_identified'] = sections
                    
                    # Analyze header/footer
                    try:
                        from docx import Document
                        doc = Document(full_path)
                        header_footer = analyze_header_footer(doc)
                        analysis['header_footer'] = header_footer
                    except:
                        pass
                    
                    all_resume_analyses.append(analysis)
                    
                    # Update stats
                    stats['total_resumes_analyzed'] += 1
                    
                    for font in analysis.get('fonts_used', []):
                        stats['fonts_frequency'][font] += 1
                    
                    for size in analysis.get('font_sizes', []):
                        stats['font_sizes_frequency'][size] += 1
                    
                    for color in analysis.get('colors_used', []):
                        stats['colors_frequency'][color] += 1
                    
                    stats['table_counts'][analysis.get('total_tables', 0)] += 1
                else:
                    print(f"  ⚠️ Failed: {analysis.get('error', 'Unknown error')}")
            
            elif file.lower().endswith('.pdf'):
                stats['pdf_count'] += 1
    
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Scan complete!")
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Total Resumés analyzed: {stats['total_resumes_analyzed']}")
    
    return all_resume_analyses, stats

def generate_detailed_report(analyses, stats, output_file):
    """Generate comprehensive detailed report about Resumé structure and formatting"""
    print(f"\n[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Generating detailed report...")
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("COMPREHENSIVE RESUMÉ ANALYSIS REPORT\n")
        f.write("Network Folder - Deep Structural Examination\n")
        f.write("=" * 80 + "\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Total Resumés Analyzed (DOCX): {stats['total_resumes_analyzed']}\n")
        f.write(f"Total Resumés Found (All): {stats['total_resumes_found']}\n")
        f.write("=" * 80 + "\n\n")
        
        # EXECUTIVE SUMMARY
        f.write("EXECUTIVE SUMMARY\n")
        f.write("-" * 80 + "\n\n")
        f.write(f"This report analyzes {stats['total_resumes_analyzed']} Resumé DOCX files.\n")
        f.write("The analysis extracts formatting, structure, tables, colors, fonts, and tone.\n")
        f.write("This provides the complete blueprint for automated Resumé generation.\n\n")
        
        # FILE FORMAT BREAKDOWN
        f.write("\n" + "=" * 80 + "\n")
        f.write("FILE FORMAT BREAKDOWN\n")
        f.write("=" * 80 + "\n\n")
        f.write(f"DOCX files (analyzed): {stats['docx_count']}\n")
        f.write(f"PDF files (counted): {stats['pdf_count']}\n")
        f.write(f"Total: {stats['total_resumes_found']}\n\n")
        
        # TYPOGRAPHY ANALYSIS
        f.write("\n" + "=" * 80 + "\n")
        f.write("TYPOGRAPHY & FONT ANALYSIS\n")
        f.write("=" * 80 + "\n\n")
        
        f.write("Font Usage (Top 20):\n")
        f.write("-" * 40 + "\n")
        for font, count in sorted(stats['fonts_frequency'].items(), key=lambda x: -x[1])[:20]:
            f.write(f"  {font:30} : {count:5} occurrences\n")
        
        f.write("\n\nFont Sizes Used (Top 15):\n")
        f.write("-" * 40 + "\n")
        for size, count in sorted(stats['font_sizes_frequency'].items(), key=lambda x: -x[1])[:15]:
            f.write(f"  {size:6.1f} pt : {count:5} occurrences\n")
        
        # COLOR ANALYSIS
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("COLOR SCHEME ANALYSIS\n")
        f.write("=" * 80 + "\n\n")
        
        if stats['colors_frequency']:
            f.write("Colors Used (RGB Values):\n")
            f.write("-" * 40 + "\n")
            for color, count in sorted(stats['colors_frequency'].items(), key=lambda x: -x[1])[:20]:
                f.write(f"  {color:20} : {count:5} times\n")
        else:
            f.write("No explicit colors detected (using default black)\n")
        
        # TABLE STRUCTURE ANALYSIS
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("TABLE STRUCTURE ANALYSIS\n")
        f.write("=" * 80 + "\n\n")
        
        f.write("Table Count Distribution:\n")
        f.write("-" * 40 + "\n")
        for table_count, resume_count in sorted(stats['table_counts'].items()):
            f.write(f"  {table_count} table(s) : {resume_count} Resumés\n")
        
        # DETAILED STRUCTURE EXAMPLES
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("DETAILED STRUCTURE EXAMPLES (First 20 Resumés)\n")
        f.write("=" * 80 + "\n\n")
        
        for i, analysis in enumerate(analyses[:20], 1):
            f.write(f"\n{i}. {analysis['person']}\n")
            f.write(f"   File: {analysis['filename']}\n")
            f.write(f"   Paragraphs: {analysis['total_paragraphs']}\n")
            f.write(f"   Tables: {analysis['total_tables']}\n")
            f.write(f"   Sections: {analysis['sections']}\n")
            
            if analysis.get('page_width'):
                f.write(f"   Page Size: {analysis['page_width']:.2f}\" × {analysis['page_height']:.2f}\"\n")
            
            if analysis.get('margin_top'):
                f.write(f"   Margins: T:{analysis['margin_top']:.2f}\" ")
                f.write(f"B:{analysis['margin_bottom']:.2f}\" ")
                f.write(f"L:{analysis['margin_left']:.2f}\" ")
                f.write(f"R:{analysis['margin_right']:.2f}\"\n")
            
            # First page content
            if analysis.get('first_page_content'):
                f.write(f"\n   First Page Content:\n")
                for j, para in enumerate(analysis['first_page_content'][:5], 1):
                    text = para['text'][:100]
                    f.write(f"     {j}. {text}\n")
                    f.write(f"        Style: {para['style']} | Align: {para['alignment']}\n")
            
            # Table structures
            if analysis.get('table_structures'):
                f.write(f"\n   Table Structures:\n")
                for table in analysis['table_structures'][:2]:  # First 2 tables
                    f.write(f"     Table {table['index']+1}: {table['rows']} rows × {table['columns']} cols\n")
                    for row_idx, row_data in enumerate(table['structure_pattern'][:3], 1):
                        f.write(f"       Row {row_idx}: {' | '.join(row_data[:3])}\n")
            
            # Sections identified
            if analysis.get('sections_identified'):
                sections = [s['section'] for s in analysis['sections_identified']]
                f.write(f"   Sections Found: {', '.join(sections)}\n")
            
            # Tone of voice
            if analysis.get('tone_of_voice'):
                tone = analysis['tone_of_voice']
                f.write(f"   Tone: Formal={tone['formal_indicators']} ")
                f.write(f"Technical={tone['technical_terms']} ")
                f.write(f"ActionVerbs={tone['action_verbs']}\n")
            
            # Header/Footer
            if analysis.get('header_footer'):
                hf = analysis['header_footer']
                if hf['has_header']:
                    f.write(f"   Header: {' | '.join(hf['header_content'][:2])}\n")
                if hf['has_footer']:
                    f.write(f"   Footer: {' | '.join(hf['footer_content'][:2])}\n")
            
            f.write("\n")
        
        # PAGE SETUP ANALYSIS
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("PAGE SETUP & MARGINS ANALYSIS\n")
        f.write("=" * 80 + "\n\n")
        
        page_sizes = [a for a in analyses if a.get('page_width')]
        if page_sizes:
            avg_width = sum(a['page_width'] for a in page_sizes) / len(page_sizes)
            avg_height = sum(a['page_height'] for a in page_sizes) / len(page_sizes)
            f.write(f"Average Page Size: {avg_width:.2f}\" × {avg_height:.2f}\" (A4 = 8.27\" × 11.69\")\n\n")
            
            avg_margin_t = sum(a['margin_top'] for a in page_sizes if a.get('margin_top')) / len(page_sizes)
            avg_margin_b = sum(a['margin_bottom'] for a in page_sizes if a.get('margin_bottom')) / len(page_sizes)
            avg_margin_l = sum(a['margin_left'] for a in page_sizes if a.get('margin_left')) / len(page_sizes)
            avg_margin_r = sum(a['margin_right'] for a in page_sizes if a.get('margin_right')) / len(page_sizes)
            
            f.write(f"Average Margins:\n")
            f.write(f"  Top: {avg_margin_t:.2f}\"\n")
            f.write(f"  Bottom: {avg_margin_b:.2f}\"\n")
            f.write(f"  Left: {avg_margin_l:.2f}\"\n")
            f.write(f"  Right: {avg_margin_r:.2f}\"\n")
        
        # TONE OF VOICE SUMMARY
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("TONE OF VOICE & WRITING STYLE\n")
        f.write("=" * 80 + "\n\n")
        
        tone_analyses = [a['tone_of_voice'] for a in analyses if a.get('tone_of_voice')]
        if tone_analyses:
            avg_formal = sum(t['formal_indicators'] for t in tone_analyses) / len(tone_analyses)
            avg_technical = sum(t['technical_terms'] for t in tone_analyses) / len(tone_analyses)
            avg_action = sum(t['action_verbs'] for t in tone_analyses) / len(tone_analyses)
            avg_bullets = sum(t['bullet_points'] for t in tone_analyses) / len(tone_analyses)
            
            f.write(f"Average per Resumé:\n")
            f.write(f"  Formal indicators: {avg_formal:.1f}\n")
            f.write(f"  Technical terms: {avg_technical:.1f}\n")
            f.write(f"  Action verbs: {avg_action:.1f}\n")
            f.write(f"  Bullet points: {avg_bullets:.1f}\n\n")
            
            f.write("Writing Style Characteristics:\n")
            f.write("  - Professional and formal tone\n")
            f.write("  - Project-focused (not personal achievements)\n")
            f.write("  - Technical terminology used\n")
            f.write("  - Action-oriented language\n")
            f.write("  - Bullet points for clarity\n")
            f.write("  - Chronological order (recent first)\n")
        
        # RECOMMENDATIONS
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("RECOMMENDATIONS FOR RESUMÉ GENERATION\n")
        f.write("=" * 80 + "\n\n")
        
        # Most common font
        if stats['fonts_frequency']:
            top_font = max(stats['fonts_frequency'].items(), key=lambda x: x[1])
            f.write(f"1. PRIMARY FONT:\n")
            f.write(f"   Use: {top_font[0]} (found in {top_font[1]} instances)\n\n")
        
        # Most common sizes
        if stats['font_sizes_frequency']:
            top_sizes = sorted(stats['font_sizes_frequency'].items(), key=lambda x: -x[1])[:3]
            f.write(f"2. FONT SIZES:\n")
            for size, count in top_sizes:
                f.write(f"   {size:.1f} pt - Used {count} times\n")
            f.write("\n")
        
        # Table usage
        if stats['table_counts']:
            most_common_tables = max(stats['table_counts'].items(), key=lambda x: x[1])
            f.write(f"3. TABLE STRUCTURE:\n")
            f.write(f"   Most common: {most_common_tables[0]} table(s) per Resumé\n")
            f.write(f"   ({most_common_tables[1]} Resumés use this pattern)\n\n")
        
        f.write(f"4. COLOR SCHEME:\n")
        if stats['colors_frequency']:
            f.write(f"   Use documented RGB colors\n")
        else:
            f.write(f"   Use default black text\n")
        f.write("\n")
        
        f.write(f"5. PAGE SETUP:\n")
        f.write(f"   A4 page size (8.27\" × 11.69\")\n")
        f.write(f"   Standard margins as analyzed\n\n")
        
        f.write(f"6. CONTENT STRUCTURE:\n")
        f.write(f"   - Header: Name, Location, Birth Year\n")
        f.write(f"   - Main section: Werkervaring (Work Experience)\n")
        f.write(f"   - Projects listed under each position\n")
        f.write(f"   - Education & Training section\n")
        f.write(f"   - Optional: Skills, Languages\n\n")
        
        f.write("=" * 80 + "\n")
        f.write("END OF REPORT\n")
        f.write("=" * 80 + "\n")
    
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Report saved to: {output_file}")

def save_json_data(analyses, output_file):
    """Save detailed analysis data as JSON"""
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Saving JSON data...")
    
    # Prepare data (keep only essential info)
    json_data = []
    for analysis in analyses:
        json_analysis = {
            'person': analysis.get('person'),
            'filename': analysis.get('filename'),
            'total_paragraphs': analysis.get('total_paragraphs'),
            'total_tables': analysis.get('total_tables'),
            'fonts_used': list(set(analysis.get('fonts_used', []))),
            'font_sizes': list(set(analysis.get('font_sizes', []))),
            'table_structures': analysis.get('table_structures', []),
            'sections_identified': analysis.get('sections_identified', []),
            'tone_of_voice': analysis.get('tone_of_voice', {}),
            'naming_pattern': analysis.get('naming_pattern', {}),
        }
        json_data.append(json_analysis)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump({
            'generated': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'total_resumes': len(analyses),
            'resume_analyses': json_data
        }, f, indent=2, ensure_ascii=False)
    
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] JSON data saved to: {output_file}")

def main():
    """Main execution function"""
    print("=" * 80)
    print("COMPREHENSIVE RESUMÉ ANALYZER")
    print("Deep Examination of All Resumés in Network Folder")
    print("=" * 80)
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Starting analysis...\n")
    
    base_path = r"C:\Users\RudoJockinSynergiepr\Synergie PM\Netwerk - Documenten"
    
    if not os.path.exists(base_path):
        print(f"ERROR: Path not found: {base_path}")
        return
    
    # Scan all Resumés
    analyses, stats = scan_all_resumes(base_path)
    
    # Generate detailed report
    report_file = "Comprehensive_Resume_Analysis_Report.txt"
    generate_detailed_report(analyses, stats, report_file)
    
    # Save JSON data
    json_file = "resume_analysis_data.json"
    save_json_data(analyses, json_file)
    
    print("\n" + "=" * 80)
    print("ANALYSIS COMPLETE!")
    print("=" * 80)
    print(f"\nFiles generated:")
    print(f"  1. {report_file} - Detailed human-readable report")
    print(f"  2. {json_file} - Machine-readable data")
    print(f"\nTotal Resumés analyzed: {stats['total_resumes_analyzed']}")
    print(f"Total Resumés found: {stats['total_resumes_found']}")
    print("\n" + "=" * 80)

if __name__ == "__main__":
    main()

