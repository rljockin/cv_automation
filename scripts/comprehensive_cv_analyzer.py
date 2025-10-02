#!/usr/bin/env python3
"""
Comprehensive CV Analyzer - Deep examination of all CVs in Network folder
Generates a detailed report for CV automation system development
"""

import os
import sys
from datetime import datetime
from pathlib import Path
import re
from collections import defaultdict, Counter
import json

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF using PyPDF2"""
    try:
        import PyPDF2
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip(), len(pdf_reader.pages)
    except Exception as e:
        return f"ERROR: {str(e)}", 0

def extract_text_from_docx(docx_path):
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(docx_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts), len(doc.paragraphs)
    except Exception as e:
        return f"ERROR: {str(e)}", 0

def detect_language(text):
    """Detect primary language of text"""
    if not text or len(text) < 50:
        return "unknown"
    
    text_lower = text.lower()
    
    # Dutch indicators
    dutch_words = ['werkervaring', 'opleiding', 'vaardigheden', 'persoonlijk', 
                   'projecten', 'ervaring', 'geboren', 'woonplaats', 'competenties',
                   'opleidingen', 'cursussen', 'talen']
    
    # English indicators
    english_words = ['work experience', 'education', 'skills', 'personal', 
                     'projects', 'experience', 'born', 'residence', 'competencies',
                     'training', 'languages', 'profile']
    
    dutch_count = sum(1 for word in dutch_words if word in text_lower)
    english_count = sum(1 for word in english_words if word in text_lower)
    
    if dutch_count > english_count:
        return "Dutch"
    elif english_count > dutch_count:
        return "English"
    elif dutch_count == english_count and dutch_count > 0:
        return "Mixed (Dutch/English)"
    else:
        return "Unknown"

def identify_sections(text):
    """Identify CV sections"""
    if not text or len(text) < 50:
        return []
    
    sections = []
    lines = text.split('\n')
    
    # Section patterns (Dutch and English)
    section_patterns = {
        'Personal Info': [r'personalia', r'personal\s+info', r'persoonlijk', r'gegevens'],
        'Profile': [r'\bprofiel\b', r'\bprofile\b', r'samenvatting', r'summary'],
        'Work Experience': [r'werkervaring', r'work\s+experience', r'ervaring', r'professional\s+experience'],
        'Education': [r'opleiding', r'education', r'opleidingen', r'training'],
        'Skills': [r'vaardigheden', r'\bskills\b', r'competenties', r'competencies'],
        'Projects': [r'projecten', r'projects', r'project\s+ervaring'],
        'Certifications': [r'certificaten', r'certificates', r'certificering', r'certifications'],
        'Languages': [r'\btalen\b', r'languages', r'talenkennis'],
        'References': [r'referenties', r'references'],
        'Courses': [r'cursussen', r'courses', r'training'],
        'Software': [r'software', r'tools', r'applicaties'],
    }
    
    for i, line in enumerate(lines):
        if len(line.strip()) > 2 and len(line.strip()) < 50:  # Likely header
            line_lower = line.lower()
            for section_name, patterns in section_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, line_lower):
                        sections.append({
                            'name': section_name,
                            'line': i + 1,
                            'text': line.strip()
                        })
                        break
    
    return sections

def detect_date_format(text):
    """Detect date formats used in CV"""
    formats = []
    
    # Various date patterns
    patterns = {
        'YYYY-MM-DD': r'\d{4}-\d{2}-\d{2}',
        'DD-MM-YYYY': r'\d{2}-\d{2}-\d{4}',
        'MM/YYYY': r'\d{2}/\d{4}',
        'YYYY': r'\b\d{4}\b',
        'Month YYYY': r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}\b',
        'YYYY - YYYY': r'\d{4}\s*[-–]\s*\d{4}',
        'YYYY - heden': r'\d{4}\s*[-–]\s*(heden|present|now)',
    }
    
    for format_name, pattern in patterns.items():
        if re.search(pattern, text, re.IGNORECASE):
            formats.append(format_name)
    
    return formats

def extract_contact_info(text):
    """Try to extract contact information patterns"""
    contact = {}
    
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    if emails:
        contact['email'] = emails[0]
    
    # Phone pattern (Dutch)
    phone_pattern = r'(\+31|0031|0)\s*[1-9]\s*\d{1,2}\s*\d{6,7}'
    phones = re.findall(phone_pattern, text.replace('-', ' ').replace('.', ' '))
    if phones:
        contact['phone'] = True
    
    # LinkedIn pattern
    if 'linkedin.com' in text.lower():
        contact['linkedin'] = True
    
    return contact

def analyze_cv_file(file_path, person_name):
    """Comprehensive analysis of a single CV file"""
    print(f"[{datetime.now().strftime('%H:%M:%S')}] Analyzing: {person_name} - {os.path.basename(file_path)}")
    
    analysis = {
        'person': person_name,
        'filename': os.path.basename(file_path),
        'filepath': file_path,
        'file_extension': os.path.splitext(file_path)[1].lower(),
        'file_size_kb': os.path.getsize(file_path) / 1024,
        'extraction_status': 'not_attempted',
        'text_length': 0,
        'page_count': 0,
        'language': 'unknown',
        'sections_found': [],
        'date_formats': [],
        'has_contact_info': {},
        'error': None
    }
    
    # Extract text
    try:
        if analysis['file_extension'] == '.pdf':
            text, pages = extract_text_from_pdf(file_path)
            analysis['page_count'] = pages
        elif analysis['file_extension'] in ['.docx', '.doc']:
            text, para_count = extract_text_from_docx(file_path)
            analysis['paragraph_count'] = para_count
        else:
            text = ""
            analysis['extraction_status'] = 'unsupported_format'
        
        if text and not text.startswith("ERROR:"):
            analysis['extraction_status'] = 'success'
            analysis['text_length'] = len(text)
            
            # Analyze content
            analysis['language'] = detect_language(text)
            analysis['sections_found'] = identify_sections(text)
            analysis['date_formats'] = detect_date_format(text)
            analysis['has_contact_info'] = extract_contact_info(text)
            
            # Store sample text (first 1000 chars)
            analysis['text_sample'] = text[:1000]
            
        elif text.startswith("ERROR:"):
            analysis['extraction_status'] = 'failed'
            analysis['error'] = text
        else:
            analysis['extraction_status'] = 'no_text_extracted'
            
    except Exception as e:
        analysis['extraction_status'] = 'exception'
        analysis['error'] = str(e)
    
    return analysis

def scan_all_cvs(base_path):
    """Scan all CV files in the Network folder"""
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Starting comprehensive CV scan...")
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Base path: {base_path}")
    
    # Install dependencies if needed
    try:
        import PyPDF2
    except ImportError:
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Installing PyPDF2...")
        os.system("pip install PyPDF2 -q")
    
    try:
        from docx import Document
    except ImportError:
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Installing python-docx...")
        os.system("pip install python-docx -q")
    
    cv_pattern = re.compile(r'^CV[_ ]', re.IGNORECASE)
    all_cv_analyses = []
    stats = {
        'total_cvs_found': 0,
        'total_cvs_analyzed': 0,
        'by_extension': defaultdict(int),
        'by_status': defaultdict(int),
        'by_language': defaultdict(int),
        'sections_frequency': defaultdict(int),
        'date_formats_used': defaultdict(int),
    }
    
    # Walk through all folders
    for root, dirs, files in os.walk(base_path):
        # Skip template folder
        if 'STANDAARD RESUME' in root:
            continue
        
        person_name = os.path.basename(root)
        
        for file in files:
            # Check if it's a CV file
            if not cv_pattern.search(file):
                continue
            
            if not file.lower().endswith(('.pdf', '.docx', '.doc')):
                continue
            
            stats['total_cvs_found'] += 1
            full_path = os.path.join(root, file)
            
            # Analyze CV
            analysis = analyze_cv_file(full_path, person_name)
            all_cv_analyses.append(analysis)
            
            # Update stats
            stats['total_cvs_analyzed'] += 1
            stats['by_extension'][analysis['file_extension']] += 1
            stats['by_status'][analysis['extraction_status']] += 1
            stats['by_language'][analysis['language']] += 1
            
            for section in analysis['sections_found']:
                stats['sections_frequency'][section['name']] += 1
            
            for date_format in analysis['date_formats']:
                stats['date_formats_used'][date_format] += 1
    
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Scan complete!")
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Total CVs analyzed: {stats['total_cvs_analyzed']}")
    
    return all_cv_analyses, stats

def generate_detailed_report(analyses, stats, output_file):
    """Generate comprehensive detailed report"""
    print(f"\n[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Generating detailed report...")
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("COMPREHENSIVE CV ANALYSIS REPORT\n")
        f.write("Network Folder - Deep Examination\n")
        f.write("=" * 80 + "\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Total CVs Analyzed: {stats['total_cvs_analyzed']}\n")
        f.write("=" * 80 + "\n\n")
        
        # EXECUTIVE SUMMARY
        f.write("EXECUTIVE SUMMARY\n")
        f.write("-" * 80 + "\n\n")
        f.write(f"This report analyzes {stats['total_cvs_analyzed']} CV files found in the Network folder.\n")
        f.write("The analysis extracts text, identifies structure, detects languages, and catalogs\n")
        f.write("all characteristics needed to build a robust CV automation system.\n\n")
        
        # STATISTICS
        f.write("\n" + "=" * 80 + "\n")
        f.write("OVERALL STATISTICS\n")
        f.write("=" * 80 + "\n\n")
        
        f.write("File Format Distribution:\n")
        f.write("-" * 40 + "\n")
        for ext, count in sorted(stats['by_extension'].items(), key=lambda x: -x[1]):
            percentage = (count / stats['total_cvs_analyzed']) * 100
            f.write(f"  {ext:10} : {count:4} files ({percentage:5.1f}%)\n")
        
        f.write("\n\nText Extraction Status:\n")
        f.write("-" * 40 + "\n")
        for status, count in sorted(stats['by_status'].items(), key=lambda x: -x[1]):
            percentage = (count / stats['total_cvs_analyzed']) * 100
            f.write(f"  {status:25} : {count:4} files ({percentage:5.1f}%)\n")
        
        f.write("\n\nLanguage Distribution:\n")
        f.write("-" * 40 + "\n")
        for lang, count in sorted(stats['by_language'].items(), key=lambda x: -x[1]):
            percentage = (count / stats['total_cvs_analyzed']) * 100
            f.write(f"  {lang:25} : {count:4} files ({percentage:5.1f}%)\n")
        
        f.write("\n\nSection Frequency (found across all CVs):\n")
        f.write("-" * 40 + "\n")
        for section, count in sorted(stats['sections_frequency'].items(), key=lambda x: -x[1]):
            percentage = (count / stats['total_cvs_analyzed']) * 100
            f.write(f"  {section:25} : {count:4} times ({percentage:5.1f}%)\n")
        
        f.write("\n\nDate Format Usage:\n")
        f.write("-" * 40 + "\n")
        for format_type, count in sorted(stats['date_formats_used'].items(), key=lambda x: -x[1]):
            percentage = (count / stats['total_cvs_analyzed']) * 100
            f.write(f"  {format_type:25} : {count:4} CVs ({percentage:5.1f}%)\n")
        
        # EXTRACTION CHALLENGES
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("EXTRACTION CHALLENGES & ISSUES\n")
        f.write("=" * 80 + "\n\n")
        
        failed = [a for a in analyses if a['extraction_status'] in ['failed', 'no_text_extracted']]
        f.write(f"Total CVs with extraction issues: {len(failed)}\n\n")
        
        if failed:
            f.write("Files requiring special handling (OCR or manual processing):\n")
            f.write("-" * 80 + "\n")
            for i, analysis in enumerate(failed[:50], 1):  # First 50
                f.write(f"{i:3}. {analysis['person']:40} | {analysis['filename']}\n")
                if analysis['error']:
                    f.write(f"     Error: {analysis['error'][:100]}\n")
        
        # LANGUAGE BREAKDOWN
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("LANGUAGE ANALYSIS\n")
        f.write("=" * 80 + "\n\n")
        
        dutch_cvs = [a for a in analyses if 'Dutch' in a['language']]
        english_cvs = [a for a in analyses if a['language'] == 'English']
        mixed_cvs = [a for a in analyses if 'Mixed' in a['language']]
        
        f.write(f"Dutch CVs: {len(dutch_cvs)}\n")
        f.write(f"English CVs: {len(english_cvs)}\n")
        f.write(f"Mixed Language CVs: {len(mixed_cvs)}\n\n")
        
        f.write("Sample English CVs (will need translation or English template):\n")
        f.write("-" * 80 + "\n")
        for i, analysis in enumerate(english_cvs[:20], 1):
            f.write(f"{i:3}. {analysis['person']:40} | {analysis['filename']}\n")
        
        # SECTION ANALYSIS
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("SECTION STRUCTURE ANALYSIS\n")
        f.write("=" * 80 + "\n\n")
        
        f.write("Common CV structures identified:\n\n")
        
        # Find CVs with most complete section coverage
        complete_cvs = []
        for analysis in analyses:
            if analysis['extraction_status'] == 'success':
                section_count = len(analysis['sections_found'])
                if section_count >= 5:
                    complete_cvs.append((section_count, analysis))
        
        complete_cvs.sort(reverse=True, key=lambda x: x[0])
        
        f.write("CVs with comprehensive section structure (5+ sections):\n")
        f.write("-" * 80 + "\n")
        for i, (section_count, analysis) in enumerate(complete_cvs[:30], 1):
            f.write(f"{i:3}. {analysis['person']:35} | Sections: {section_count} | {analysis['filename']}\n")
            sections = [s['name'] for s in analysis['sections_found']]
            f.write(f"     Sections found: {', '.join(sections)}\n")
        
        # FILE SIZE ANALYSIS
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("FILE SIZE ANALYSIS\n")
        f.write("=" * 80 + "\n\n")
        
        sizes = [a['file_size_kb'] for a in analyses]
        avg_size = sum(sizes) / len(sizes) if sizes else 0
        max_size = max(sizes) if sizes else 0
        min_size = min(sizes) if sizes else 0
        
        f.write(f"Average CV file size: {avg_size:.1f} KB\n")
        f.write(f"Largest CV: {max_size:.1f} KB\n")
        f.write(f"Smallest CV: {min_size:.1f} KB\n\n")
        
        # Large files (might be image-heavy)
        large_files = [a for a in analyses if a['file_size_kb'] > 1000]  # > 1MB
        f.write(f"\nLarge files (>1MB, possibly image-heavy): {len(large_files)}\n")
        f.write("-" * 80 + "\n")
        for i, analysis in enumerate(large_files[:20], 1):
            f.write(f"{i:3}. {analysis['person']:35} | {analysis['file_size_kb']:.0f} KB | {analysis['filename']}\n")
        
        # DETAILED CV CATALOG
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("DETAILED CV CATALOG (All CVs)\n")
        f.write("=" * 80 + "\n\n")
        
        # Group by extraction status
        success = [a for a in analyses if a['extraction_status'] == 'success']
        
        f.write(f"\n\nSuccessfully Extracted CVs ({len(success)}):\n")
        f.write("=" * 80 + "\n\n")
        
        for i, analysis in enumerate(success, 1):
            f.write(f"{i:4}. {analysis['person']}\n")
            f.write(f"      File: {analysis['filename']}\n")
            f.write(f"      Format: {analysis['file_extension']} | Size: {analysis['file_size_kb']:.1f} KB\n")
            f.write(f"      Language: {analysis['language']} | Text Length: {analysis['text_length']} chars\n")
            if analysis.get('page_count'):
                f.write(f"      Pages: {analysis['page_count']}\n")
            
            if analysis['sections_found']:
                sections = [s['name'] for s in analysis['sections_found']]
                f.write(f"      Sections: {', '.join(sections)}\n")
            
            if analysis['date_formats']:
                f.write(f"      Date Formats: {', '.join(analysis['date_formats'])}\n")
            
            if analysis['has_contact_info']:
                contact = []
                if analysis['has_contact_info'].get('email'):
                    contact.append('Email')
                if analysis['has_contact_info'].get('phone'):
                    contact.append('Phone')
                if analysis['has_contact_info'].get('linkedin'):
                    contact.append('LinkedIn')
                if contact:
                    f.write(f"      Contact Info: {', '.join(contact)}\n")
            
            f.write("\n")
        
        # RECOMMENDATIONS
        f.write("\n\n" + "=" * 80 + "\n")
        f.write("RECOMMENDATIONS FOR CV AUTOMATION SYSTEM\n")
        f.write("=" * 80 + "\n\n")
        
        failed_percentage = (len(failed) / stats['total_cvs_analyzed']) * 100 if stats['total_cvs_analyzed'] > 0 else 0
        
        f.write("1. OCR REQUIREMENT:\n")
        f.write(f"   - {len(failed)} CVs ({failed_percentage:.1f}%) require OCR processing\n")
        f.write("   - Implement Tesseract OCR for image-based PDFs\n")
        f.write("   - Consider pdf2image + pytesseract pipeline\n\n")
        
        f.write("2. MULTI-LANGUAGE SUPPORT:\n")
        f.write(f"   - Dutch CVs: {len(dutch_cvs)}\n")
        f.write(f"   - English CVs: {len(english_cvs)}\n")
        f.write("   - Need language detection and appropriate template mapping\n\n")
        
        f.write("3. SECTION IDENTIFICATION:\n")
        f.write("   - Implement flexible section detection (regex patterns)\n")
        f.write("   - Handle both Dutch and English section names\n")
        f.write("   - Most common sections to extract:\n")
        for section, count in sorted(stats['sections_frequency'].items(), key=lambda x: -x[1])[:10]:
            f.write(f"     * {section}: found in {count} CVs\n")
        f.write("\n")
        
        f.write("4. DATE PARSING:\n")
        f.write("   - Support multiple date formats:\n")
        for format_type in sorted(stats['date_formats_used'].keys()):
            f.write(f"     * {format_type}\n")
        f.write("\n")
        
        f.write("5. FORMAT HANDLING:\n")
        for ext, count in sorted(stats['by_extension'].items(), key=lambda x: -x[1]):
            f.write(f"   - {ext}: {count} files - ")
            if ext == '.pdf':
                f.write("Use PyPDF2 with OCR fallback\n")
            elif ext in ['.docx', '.doc']:
                f.write("Use python-docx library\n")
            else:
                f.write("Consider manual processing\n")
        
        f.write("\n6. QUALITY CONTROL:\n")
        f.write(f"   - Implement validation for extracted data\n")
        f.write(f"   - Flag CVs with < 500 characters for manual review\n")
        f.write(f"   - Create review queue for {len(failed)} problematic CVs\n\n")
        
        f.write("=" * 80 + "\n")
        f.write("END OF REPORT\n")
        f.write("=" * 80 + "\n")
    
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Report saved to: {output_file}")

def save_json_data(analyses, output_file):
    """Save detailed analysis data as JSON for programmatic access"""
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Saving JSON data...")
    
    # Prepare data (remove text samples to keep file size manageable)
    json_data = []
    for analysis in analyses:
        analysis_copy = analysis.copy()
        if 'text_sample' in analysis_copy:
            del analysis_copy['text_sample']
        json_data.append(analysis_copy)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump({
            'generated': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'total_cvs': len(analyses),
            'cv_analyses': json_data
        }, f, indent=2, ensure_ascii=False)
    
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] JSON data saved to: {output_file}")

def main():
    """Main execution function"""
    print("=" * 80)
    print("COMPREHENSIVE CV ANALYZER")
    print("Deep Examination of All CVs in Network Folder")
    print("=" * 80)
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Starting analysis...\n")
    
    base_path = r"C:\Users\RudoJockinSynergiepr\Synergie PM\Netwerk - Documenten"
    
    if not os.path.exists(base_path):
        print(f"ERROR: Path not found: {base_path}")
        return
    
    # Scan all CVs
    analyses, stats = scan_all_cvs(base_path)
    
    # Generate detailed report
    report_file = "Comprehensive_CV_Analysis_Report.txt"
    generate_detailed_report(analyses, stats, report_file)
    
    # Save JSON data
    json_file = "cv_analysis_data.json"
    save_json_data(analyses, json_file)
    
    print("\n" + "=" * 80)
    print("ANALYSIS COMPLETE!")
    print("=" * 80)
    print(f"\nFiles generated:")
    print(f"  1. {report_file} - Detailed human-readable report")
    print(f"  2. {json_file} - Machine-readable data for automation")
    print(f"\nTotal CVs analyzed: {stats['total_cvs_analyzed']}")
    print(f"Successfully extracted: {stats['by_status']['success']}")
    print(f"Requiring OCR: {stats['by_status']['failed'] + stats['by_status']['no_text_extracted']}")
    print("\n" + "=" * 80)

if __name__ == "__main__":
    main()

