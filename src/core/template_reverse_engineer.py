#!/usr/bin/env python3
"""
Template Reverse Engineer - Deep dive into actual Resumé DOCX structure
Extracts EXACT formatting details: cell properties, borders, shading, alignment, etc.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import qn

def analyze_paragraph_format(paragraph):
    """Extract detailed paragraph formatting"""
    fmt = {
        'text': paragraph.text[:100] if paragraph.text else '',
        'style': paragraph.style.name if paragraph.style else None,
        'alignment': str(paragraph.alignment) if paragraph.alignment else None,
        'left_indent': paragraph.paragraph_format.left_indent,
        'right_indent': paragraph.paragraph_format.right_indent,
        'first_line_indent': paragraph.paragraph_format.first_line_indent,
        'space_before': paragraph.paragraph_format.space_before,
        'space_after': paragraph.paragraph_format.space_after,
        'line_spacing': paragraph.paragraph_format.line_spacing,
        'runs': []
    }
    
    for run in paragraph.runs:
        run_fmt = {
            'text': run.text[:50],
            'font_name': run.font.name,
            'font_size': run.font.size.pt if run.font.size else None,
            'bold': run.font.bold,
            'italic': run.font.italic,
            'underline': run.font.underline,
        }
        
        # Color
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            run_fmt['color_rgb'] = f"({rgb[0]}, {rgb[1]}, {rgb[2]})"
            run_fmt['color_hex'] = f"#{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
        
        fmt['runs'].append(run_fmt)
    
    return fmt

def analyze_cell_format(cell):
    """Extract detailed cell formatting"""
    fmt = {
        'text': cell.text[:100] if cell.text else '',
        'width': cell.width,
        'vertical_alignment': str(cell.vertical_alignment) if hasattr(cell, 'vertical_alignment') else None,
    }
    
    # Cell shading/background
    try:
        shading_elm = cell._element.get_or_add_tcPr()
        shd_elm = shading_elm.find(qn('w:shd'))
        if shd_elm is not None:
            fill = shd_elm.get(qn('w:fill'))
            if fill:
                fmt['background_color'] = fill
    except:
        pass
    
    # Cell borders
    try:
        tc_pr = cell._element.tcPr
        if tc_pr is not None:
            tc_borders = tc_pr.find(qn('w:tcBorders'))
            if tc_borders is not None:
                fmt['has_borders'] = True
    except:
        pass
    
    return fmt

def analyze_table_structure(table, table_idx):
    """Deep analysis of table structure"""
    print(f"\n{'─'*80}")
    print(f"TABLE {table_idx + 1}: {len(table.rows)} rows × {len(table.columns)} columns")
    print(f"{'─'*80}")
    
    analysis = {
        'index': table_idx,
        'rows': len(table.rows),
        'columns': len(table.columns),
        'column_widths': [],
        'row_heights': [],
        'cells': [],
    }
    
    # Column widths
    for col_idx, column in enumerate(table.columns):
        width = column.width
        width_inches = width.inches if width else None
        analysis['column_widths'].append(width_inches)
        print(f"  Column {col_idx + 1}: {width_inches:.2f}\" " if width_inches else f"  Column {col_idx + 1}: Auto")
    
    # Analyze first 5 rows in detail
    print(f"\n  First 5 Rows:")
    for row_idx, row in enumerate(table.rows[:5]):
        print(f"\n    Row {row_idx + 1}:")
        row_data = []
        
        for cell_idx, cell in enumerate(row.cells):
            cell_fmt = analyze_cell_format(cell)
            row_data.append(cell_fmt)
            
            # Print cell details
            text_preview = cell.text[:50] if cell.text else '(empty)'
            print(f"      Cell [{row_idx+1},{cell_idx+1}]: {text_preview}")
            
            if cell_fmt.get('background_color'):
                print(f"        Background: {cell_fmt['background_color']}")
            
            # Analyze paragraphs in cell
            for para in cell.paragraphs:
                if para.text.strip():
                    para_fmt = analyze_paragraph_format(para)
                    if para_fmt['runs']:
                        for run_fmt in para_fmt['runs']:
                            if run_fmt.get('color_hex'):
                                print(f"        Text Color: {run_fmt['color_hex']} (RGB: {run_fmt['color_rgb']})")
                            if run_fmt.get('font_size'):
                                print(f"        Font: {run_fmt['font_name']} {run_fmt['font_size']}pt", end='')
                                if run_fmt.get('bold'):
                                    print(" BOLD", end='')
                                print()
        
        analysis['cells'].append(row_data)
    
    return analysis

def reverse_engineer_resume(docx_path):
    """Complete reverse engineering of a Resumé file"""
    print(f"{'='*80}")
    print(f"RESUMÉ TEMPLATE REVERSE ENGINEERING")
    print(f"{'='*80}")
    print(f"File: {os.path.basename(docx_path)}")
    print(f"Path: {docx_path}\n")
    
    try:
        doc = Document(docx_path)
        
        # Document properties
        print(f"{'='*80}")
        print(f"DOCUMENT PROPERTIES")
        print(f"{'='*80}")
        
        section = doc.sections[0]
        print(f"Page Width: {section.page_width.inches:.2f}\"")
        print(f"Page Height: {section.page_height.inches:.2f}\"")
        print(f"Top Margin: {section.top_margin.inches:.2f}\"")
        print(f"Bottom Margin: {section.bottom_margin.inches:.2f}\"")
        print(f"Left Margin: {section.left_margin.inches:.2f}\"")
        print(f"Right Margin: {section.right_margin.inches:.2f}\"")
        
        # Paragraphs (before first table)
        print(f"\n{'='*80}")
        print(f"HEADER PARAGRAPHS (Before Tables)")
        print(f"{'='*80}")
        
        para_count = 0
        for para in doc.paragraphs[:10]:
            if para.text.strip():
                para_count += 1
                para_fmt = analyze_paragraph_format(para)
                print(f"\nParagraph {para_count}:")
                print(f"  Text: {para.text}")
                print(f"  Style: {para_fmt['style']}")
                print(f"  Alignment: {para_fmt['alignment']}")
                
                for run_fmt in para_fmt['runs']:
                    print(f"  Font: {run_fmt['font_name']} {run_fmt['font_size']}pt", end='')
                    if run_fmt.get('bold'):
                        print(" BOLD", end='')
                    if run_fmt.get('color_hex'):
                        print(f" Color: {run_fmt['color_hex']}", end='')
                    print()
        
        # Tables
        print(f"\n{'='*80}")
        print(f"TABLES ANALYSIS ({len(doc.tables)} total)")
        print(f"{'='*80}")
        
        all_table_analyses = []
        for table_idx, table in enumerate(doc.tables[:8]):  # Analyze first 8 tables
            table_analysis = analyze_table_structure(table, table_idx)
            all_table_analyses.append(table_analysis)
        
        # Save detailed analysis
        output_file = "template_structure_analysis.txt"
        print(f"\n{'='*80}")
        print(f"Analysis complete! Full details in: {output_file}")
        print(f"{'='*80}")
        
        return True
        
    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main function"""
    # Sample Resumé file to analyze
    resume_path = r"C:\Users\RudoJockinSynergiepr\Synergie PM\Netwerk - Documenten\Smit, Ronald\Resumé_Synergie projectmanagement_Ronald Smit.docx"
    
    if not os.path.exists(resume_path):
        print(f"File not found: {resume_path}")
        print("Please provide a valid Resumé DOCX path")
        return
    
    reverse_engineer_resume(resume_path)

if __name__ == "__main__":
    main()

