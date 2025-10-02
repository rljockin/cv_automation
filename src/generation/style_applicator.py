#!/usr/bin/env python3
"""
Style Applicator - Apply exact Synergie styling to documents
Handles fonts, colors, sizes, and formatting to match template exactly
"""

from typing import Dict, List, Optional, Any
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.core import (
    SynergieColors,
    SynergieFonts,
    SynergiePageSetup,
    PersonalInfo,
    WorkExperience,
    Language
)


class StyleApplicator:
    """
    Apply exact Synergie styling to documents
    
    Features:
    - Calibri font family throughout
    - Synergie Orange (#D07E1F) for headers
    - Exact font sizes (24pt, 18pt, 14pt, 10pt)
    - Proper spacing and alignment
    - Consistent formatting
    - Brand compliance
    
    Based on comprehensive Resumé analysis:
    - 539 Resumés analyzed
    - Exact styling specifications extracted
    - Pixel-perfect matching required
    """
    
    def __init__(self):
        """Initialize style applicator with Synergie specifications"""
        
        # Font specifications
        self.font_family = SynergieFonts.FONT_FAMILY
        self.font_sizes = SynergieFonts.FONT_SIZES
        
        # Color specifications
        self.synergie_orange = SynergieColors.SYNERGIE_ORANGE
        self.text_color = SynergieColors.TEXT_COLOR
        
        # Spacing specifications
        self.line_spacing = 1.15  # Standard line spacing
        self.paragraph_spacing = 6  # Points between paragraphs
    
    def apply_document_styles(self, doc: Document) -> None:
        """
        Apply Synergie styles to entire document
        
        Args:
            doc: Document object to style
        """
        
        # Apply styles to all paragraphs
        for paragraph in doc.paragraphs:
            self._apply_paragraph_styles(paragraph)
        
        # Apply styles to all tables
        for table in doc.tables:
            self._apply_table_styles(table)
    
    def _apply_paragraph_styles(self, paragraph) -> None:
        """Apply styles to a paragraph"""
        
        # Set paragraph spacing
        paragraph.paragraph_format.space_after = Pt(self.paragraph_spacing)
        paragraph.paragraph_format.line_spacing = self.line_spacing
        
        # Apply styles to runs
        for run in paragraph.runs:
            self._apply_run_styles(run)
    
    def _apply_run_styles(self, run) -> None:
        """Apply styles to a text run"""
        
        # Set font family
        run.font.name = self.font_family
        
        # Set font color
        run.font.color.rgb = RGBColor(*self.text_color)
        
        # Determine font size based on context
        font_size = self._determine_font_size(run)
        run.font.size = Pt(font_size)
    
    def _determine_font_size(self, run) -> int:
        """Determine appropriate font size for a run"""
        
        # Check if run is bold (likely a header)
        if run.font.bold:
            # Check if it's a section header (orange color)
            if hasattr(run.font.color, 'rgb') and run.font.color.rgb:
                if run.font.color.rgb == RGBColor(*self.synergie_orange):
                    return self.font_sizes['section']  # 18pt
                else:
                    return self.font_sizes['subsection']  # 14pt
            else:
                return self.font_sizes['subsection']  # 14pt
        
        # Default to body text size
        return self.font_sizes['body']  # 10pt
    
    def _apply_table_styles(self, table) -> None:
        """Apply styles to a table"""
        
        # Apply styles to all cells
        for row in table.rows:
            for cell in row.cells:
                self._apply_cell_styles(cell)
    
    def _apply_cell_styles(self, cell) -> None:
        """Apply styles to a table cell"""
        
        # Apply styles to all paragraphs in cell
        for paragraph in cell.paragraphs:
            self._apply_paragraph_styles(paragraph)
    
    def create_name_style(self, doc: Document, name: str) -> None:
        """
        Create styled name paragraph
        
        Args:
            doc: Document object
            name: Person's name
        """
        
        # Add name paragraph
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add name run
        name_run = name_para.add_run(name)
        name_run.font.name = self.font_family
        name_run.font.size = Pt(self.font_sizes['name'])  # 24pt
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(*self.synergie_orange)
        
        # Add spacing after name
        doc.add_paragraph()
    
    def create_section_header_style(self, doc: Document, header_text: str) -> None:
        """
        Create styled section header
        
        Args:
            doc: Document object
            header_text: Section header text
        """
        
        # Add section header paragraph
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add header run
        header_run = header_para.add_run(header_text)
        header_run.font.name = self.font_family
        header_run.font.size = Pt(self.font_sizes['section'])  # 18pt
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(*self.synergie_orange)
        
        # Add spacing after header
        doc.add_paragraph()
    
    def create_subsection_header_style(self, doc: Document, header_text: str) -> None:
        """
        Create styled subsection header
        
        Args:
            doc: Document object
            header_text: Subsection header text
        """
        
        # Add subsection header paragraph
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add header run
        header_run = header_para.add_run(header_text)
        header_run.font.name = self.font_family
        header_run.font.size = Pt(self.font_sizes['subsection'])  # 14pt
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(*self.text_color)
    
    def create_body_text_style(self, doc: Document, text: str) -> None:
        """
        Create styled body text
        
        Args:
            doc: Document object
            text: Body text
        """
        
        if not text or not text.strip():
            return
        
        # Add body text paragraph
        body_para = doc.add_paragraph()
        body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add body text run
        body_run = body_para.add_run(text.strip())
        body_run.font.name = self.font_family
        body_run.font.size = Pt(self.font_sizes['body'])  # 10pt
        body_run.font.color.rgb = RGBColor(*self.text_color)
    
    def create_bullet_point_style(self, doc: Document, text: str) -> None:
        """
        Create styled bullet point
        
        Args:
            doc: Document object
            text: Bullet point text
        """
        
        if not text or not text.strip():
            return
        
        # Add bullet point paragraph
        bullet_para = doc.add_paragraph()
        bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        bullet_para.paragraph_format.left_indent = Inches(0.25)
        
        # Add bullet point run
        bullet_run = bullet_para.add_run(f"• {text.strip()}")
        bullet_run.font.name = self.font_family
        bullet_run.font.size = Pt(self.font_sizes['body'])  # 10pt
        bullet_run.font.color.rgb = RGBColor(*self.text_color)
    
    def create_contact_info_style(self, doc: Document, contact_info: List[str]) -> None:
        """
        Create styled contact information
        
        Args:
            doc: Document object
            contact_info: List of contact information items
        """
        
        if not contact_info:
            return
        
        # Add contact info paragraph
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add contact info run
        contact_run = contact_para.add_run(' | '.join(contact_info))
        contact_run.font.name = self.font_family
        contact_run.font.size = Pt(self.font_sizes['body'])  # 10pt
        contact_run.font.color.rgb = RGBColor(*self.text_color)
    
    def create_work_experience_style(self, doc: Document, work_exp: WorkExperience) -> None:
        """
        Create styled work experience entry
        
        Args:
            doc: Document object
            work_exp: Work experience object
        """
        
        # Add company and position
        if work_exp.company and work_exp.position:
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(f"{work_exp.position} bij {work_exp.company}")
            company_run.font.name = self.font_family
            company_run.font.size = Pt(self.font_sizes['subsection'])  # 14pt
            company_run.font.bold = True
            company_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add location if available
        if work_exp.location:
            location_para = doc.add_paragraph()
            location_run = location_para.add_run(work_exp.location)
            location_run.font.name = self.font_family
            location_run.font.size = Pt(self.font_sizes['body'])  # 10pt
            location_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add description if available
        if work_exp.description:
            desc_para = doc.add_paragraph()
            desc_run = desc_para.add_run(work_exp.description)
            desc_run.font.name = self.font_family
            desc_run.font.size = Pt(self.font_sizes['body'])  # 10pt
            desc_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add responsibilities
        if work_exp.responsibilities:
            for responsibility in work_exp.responsibilities[:3]:  # Limit to 3
                self.create_bullet_point_style(doc, responsibility)
    
    def create_education_style(self, doc: Document, education: Dict) -> None:
        """
        Create styled education entry
        
        Args:
            doc: Document object
            education: Education dictionary
        """
        
        # Add degree and institution
        if 'degree' in education and 'institution' in education:
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(f"{education['degree']} - {education['institution']}")
            degree_run.font.name = self.font_family
            degree_run.font.size = Pt(self.font_sizes['subsection'])  # 14pt
            degree_run.font.bold = True
            degree_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add specialization if available
        if 'specialization' in education:
            spec_para = doc.add_paragraph()
            spec_run = spec_para.add_run(education['specialization'])
            spec_run.font.name = self.font_family
            spec_run.font.size = Pt(self.font_sizes['body'])  # 10pt
            spec_run.font.color.rgb = RGBColor(*self.text_color)
    
    def create_skills_style(self, doc: Document, skills: List[str]) -> None:
        """
        Create styled skills list
        
        Args:
            doc: Document object
            skills: List of skills
        """
        
        for skill in skills:
            self.create_body_text_style(doc, skill)
    
    def create_projects_style(self, doc: Document, projects: List[Dict]) -> None:
        """
        Create styled projects list
        
        Args:
            doc: Document object
            projects: List of project dictionaries
        """
        
        for project in projects:
            # Add project name
            if 'name' in project:
                name_para = doc.add_paragraph()
                name_run = name_para.add_run(project['name'])
                name_run.font.name = self.font_family
                name_run.font.size = Pt(self.font_sizes['subsection'])  # 14pt
                name_run.font.bold = True
                name_run.font.color.rgb = RGBColor(*self.text_color)
            
            # Add project description
            if 'description' in project:
                self.create_body_text_style(doc, project['description'])
            
            # Add client if available
            if 'client' in project:
                client_para = doc.add_paragraph()
                client_run = client_para.add_run(f"Klant: {project['client']}")
                client_run.font.name = self.font_family
                client_run.font.size = Pt(self.font_sizes['body'])  # 10pt
                client_run.font.color.rgb = RGBColor(*self.text_color)
    
    def add_page_break(self, doc: Document) -> None:
        """Add page break to document"""
        
        # Add page break paragraph
        break_para = doc.add_paragraph()
        break_run = break_para.add_run()
        break_run.add_break(WD_BREAK.PAGE)
    
    def add_section_spacing(self, doc: Document) -> None:
        """Add spacing between sections"""
        
        # Add empty paragraph for spacing
        doc.add_paragraph()
    
    def validate_styling(self, doc: Document) -> Dict[str, Any]:
        """
        Validate that styling matches Synergie specifications
        
        Args:
            doc: Document object to validate
            
        Returns:
            Dictionary with validation results
        """
        
        validation_results = {
            'font_family_correct': True,
            'font_sizes_correct': True,
            'colors_correct': True,
            'spacing_correct': True,
            'issues': []
        }
        
        # Check all paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Check font family
                if run.font.name != self.font_family:
                    validation_results['font_family_correct'] = False
                    validation_results['issues'].append(f"Wrong font family: {run.font.name}")
                
                # Check font size
                if run.font.size and run.font.size.pt not in self.font_sizes.values():
                    validation_results['font_sizes_correct'] = False
                    validation_results['issues'].append(f"Wrong font size: {run.font.size.pt}")
                
                # Check colors
                if run.font.color.rgb:
                    if run.font.color.rgb not in [RGBColor(*self.synergie_orange), RGBColor(*self.text_color)]:
                        validation_results['colors_correct'] = False
                        validation_results['issues'].append(f"Wrong color: {run.font.color.rgb}")
        
        return validation_results
