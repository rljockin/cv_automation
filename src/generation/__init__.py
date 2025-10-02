#!/usr/bin/env python3
"""
Resume Generator - Generate Synergie Resumés from structured data
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.core.constants import SynergieColors, SynergieFonts, FontSizes

class ResumeGenerator:
    """Generate Synergie Resumés from CV data"""
    
    def __init__(self):
        pass
    
    def generate_resume(self, cv_data_dict, output_path, convert_to_pdf=False):
        """Generate Resumé from CV data"""
        try:
            # Create a new DOCX document
            doc = Document()
            
            # Set up page layout
            self._setup_page_layout(doc)
            
            # Add personal information
            personal_info = cv_data_dict.get('personal_info', {})
            name = personal_info.get('full_name', 'Unknown')
            location = personal_info.get('location', 'Unknown Location')
            
            # Add name with proper styling
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            name_run = name_para.add_run(name)
            name_run.font.name = SynergieFonts.PRIMARY
            name_run.font.size = FontSizes.NAME_LARGE
            name_run.font.bold = True
            name_run.font.color.rgb = SynergieColors.SECTION_HEADER
            
            # Add location
            location_para = doc.add_paragraph()
            location_run = location_para.add_run(location)
            location_run.font.name = SynergieFonts.PRIMARY
            location_run.font.size = FontSizes.BODY_TEXT
            location_run.font.color.rgb = SynergieColors.BLACK
            
            # Add spacing
            doc.add_paragraph()
            
            # Add work experience section
            work_experience = cv_data_dict.get('work_experience', [])
            if work_experience:
                self._add_section_header(doc, "Werkervaring")
                
                for exp in work_experience[:3]:  # Limit to 3 most recent
                    if isinstance(exp, dict):
                        company = exp.get('company', 'Unknown Company')
                        position = exp.get('position', 'Unknown Position')
                        
                        # Add company and position
                        exp_para = doc.add_paragraph()
                        exp_run = exp_para.add_run(f"{position} bij {company}")
                        exp_run.font.name = SynergieFonts.PRIMARY
                        exp_run.font.size = FontSizes.SUB_HEADER
                        exp_run.font.bold = True
                        exp_run.font.color.rgb = SynergieColors.BLACK
                        
                        # Add description if available
                        description = exp.get('description', '')
                        if description:
                            desc_para = doc.add_paragraph()
                            desc_run = desc_para.add_run(description[:200] + "..." if len(description) > 200 else description)
                            desc_run.font.name = SynergieFonts.PRIMARY
                            desc_run.font.size = FontSizes.BODY_TEXT
                            desc_run.font.color.rgb = SynergieColors.BLACK
                        
                        doc.add_paragraph()  # Add spacing
            
            # Add education section
            education = cv_data_dict.get('education', [])
            if education:
                self._add_section_header(doc, "Opleiding")
                
                for edu in education[:2]:  # Limit to 2 most recent
                    if isinstance(edu, dict):
                        degree = edu.get('degree', 'Unknown Degree')
                        institution = edu.get('institution', 'Unknown Institution')
                        
                        edu_para = doc.add_paragraph()
                        edu_run = edu_para.add_run(f"{degree} - {institution}")
                        edu_run.font.name = SynergieFonts.PRIMARY
                        edu_run.font.size = FontSizes.SUB_HEADER
                        edu_run.font.bold = True
                        edu_run.font.color.rgb = SynergieColors.BLACK
                        
                        doc.add_paragraph()  # Add spacing
            
            # Add skills section if available
            skills = cv_data_dict.get('skills', [])
            if skills:
                self._add_section_header(doc, "Vaardigheden")
                
                skills_text = " | ".join(skills[:10])  # Limit to 10 skills
                skills_para = doc.add_paragraph()
                skills_run = skills_para.add_run(skills_text)
                skills_run.font.name = SynergieFonts.PRIMARY
                skills_run.font.size = FontSizes.BODY_TEXT
                skills_run.font.color.rgb = SynergieColors.BLACK
            
            # Save document
            doc.save(output_path)
            
            return {
                'success': True,
                'docx_path': output_path,
                'pdf_path': None,
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'docx_path': None,
                'pdf_path': None,
                'error': str(e)
            }
    
    def _setup_page_layout(self, doc):
        """Set up page layout with Synergie specifications"""
        section = doc.sections[0]
        
        # Set page size to A4
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        
        # Set margins
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    def _add_section_header(self, doc, header_text):
        """Add a section header with proper styling"""
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_para.add_run(header_text)
        header_run.font.name = SynergieFonts.PRIMARY
        header_run.font.size = FontSizes.SECTION_HEADER
        header_run.font.bold = True
        header_run.font.color.rgb = SynergieColors.SECTION_HEADER
        
        # Add spacing after header
        doc.add_paragraph()

__all__ = ['ResumeGenerator']