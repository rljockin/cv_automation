#!/usr/bin/env python3
"""
Template Engine - Create DOCX documents with exact Synergie template
Sets up page layout, margins, fonts, and document structure
"""

import os
from typing import Optional, Dict, Any
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn

from src.core import (
    ProcessingConfig,
    SynergieColors,
    SynergieFonts,
    SynergiePageSetup,
    SynergieTableSpecs,
    CVFile,
    PersonalInfo,
    WorkExperience,
    Language
)


class TemplateEngine:
    """
    Create DOCX documents with exact Synergie template specifications
    
    Features:
    - A4 page size with exact margins
    - Calibri font family
    - Synergie Orange (#D07E1F) branding
    - Proper font sizes (24pt, 18pt, 14pt, 10pt)
    - Document structure setup
    - Header and footer configuration
    
    Based on comprehensive Resumé analysis:
    - 539 Resumés analyzed
    - Exact template specifications extracted
    - Pixel-perfect matching required
    """
    
    def __init__(self):
        """Initialize template engine with Synergie specifications"""
        
        # Document specifications from analysis
        self.page_width = SynergiePageSetup.PAGE_WIDTH
        self.page_height = SynergiePageSetup.PAGE_HEIGHT
        self.margin_top = SynergiePageSetup.MARGIN_TOP
        self.margin_bottom = SynergiePageSetup.MARGIN_BOTTOM
        self.margin_left = SynergiePageSetup.MARGIN_LEFT
        self.margin_right = SynergiePageSetup.MARGIN_RIGHT
        
        # Font specifications
        self.font_family = SynergieFonts.FONT_FAMILY
        self.font_sizes = SynergieFonts.FONT_SIZES
        
        # Color specifications
        self.synergie_orange = SynergieColors.SYNERGIE_ORANGE
        self.text_color = SynergieColors.TEXT_COLOR
        
        # Table specifications
        self.table_specs = SynergieTableSpecs
    
    def create_document(self, personal_info: PersonalInfo) -> Document:
        """
        Create new DOCX document with Synergie template
        
        Args:
            personal_info: Personal information for document setup
            
        Returns:
            Document object with Synergie template applied
        """
        
        # Create new document
        doc = Document()
        
        # Set up page layout
        self._setup_page_layout(doc)
        
        # Set up fonts and styles
        self._setup_styles(doc)
        
        # Add document header
        self._add_document_header(doc, personal_info)
        
        return doc
    
    def _setup_page_layout(self, doc: Document) -> None:
        """Set up page layout with exact Synergie specifications"""
        
        # Get first section
        section = doc.sections[0]
        
        # Set page size to A4
        section.page_width = Inches(self.page_width)
        section.page_height = Inches(self.page_height)
        
        # Set margins (exact measurements from analysis)
        section.top_margin = Inches(self.margin_top)
        section.bottom_margin = Inches(self.margin_bottom)
        section.left_margin = Inches(self.margin_left)
        section.right_margin = Inches(self.margin_right)
        
        # Set page orientation (portrait) - handled by page dimensions
    
    def _setup_styles(self, doc: Document) -> None:
        """Set up document styles with Synergie specifications"""
        
        # Get styles object
        styles = doc.styles
        
        # Set default font for document
        self._set_default_font(doc)
        
        # Create custom styles if they don't exist
        self._create_custom_styles(styles)
    
    def _set_default_font(self, doc: Document) -> None:
        """Set default font for the document"""
        
        # Set default paragraph style
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_family
        font.size = Pt(self.font_sizes['body'])
        font.color.rgb = RGBColor(*self.text_color)
    
    def _create_custom_styles(self, styles) -> None:
        """Create custom styles for different text elements"""
        
        # Name style (24pt, bold, orange)
        if 'Synergie Name' not in [s.name for s in styles]:
            name_style = styles.add_style('Synergie Name', 1)  # Paragraph style
            name_font = name_style.font
            name_font.name = self.font_family
            name_font.size = Pt(self.font_sizes['name'])
            name_font.bold = True
            name_font.color.rgb = RGBColor(*self.synergie_orange)
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Section header style (18pt, bold, orange)
        if 'Synergie Section' not in [s.name for s in styles]:
            section_style = styles.add_style('Synergie Section', 1)
            section_font = section_style.font
            section_font.name = self.font_family
            section_font.size = Pt(self.font_sizes['section'])
            section_font.bold = True
            section_font.color.rgb = RGBColor(*self.synergie_orange)
            section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Subsection style (14pt, bold, black)
        if 'Synergie Subsection' not in [s.name for s in styles]:
            subsection_style = styles.add_style('Synergie Subsection', 1)
            subsection_font = subsection_style.font
            subsection_font.name = self.font_family
            subsection_font.size = Pt(self.font_sizes['subsection'])
            subsection_font.bold = True
            subsection_font.color.rgb = RGBColor(*self.text_color)
            subsection_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Body text style (10pt, normal, black)
        if 'Synergie Body' not in [s.name for s in styles]:
            body_style = styles.add_style('Synergie Body', 1)
            body_font = body_style.font
            body_font.name = self.font_family
            body_font.size = Pt(self.font_sizes['body'])
            body_font.color.rgb = RGBColor(*self.text_color)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_document_header(self, doc: Document, personal_info: PersonalInfo) -> None:
        """Add document header with personal information"""
        
        # Add name
        if personal_info.full_name:
            name_para = doc.add_paragraph()
            name_para.style = 'Synergie Name'
            name_run = name_para.add_run(personal_info.full_name)
            name_run.font.name = self.font_family
            name_run.font.size = Pt(self.font_sizes['name'])
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(*self.synergie_orange)
        
        # Add contact information
        contact_info = []
        
        if personal_info.location:
            contact_info.append(personal_info.location)
        
        if personal_info.phone:
            contact_info.append(personal_info.phone)
        
        if personal_info.email:
            contact_info.append(personal_info.email)
        
        if personal_info.website:
            contact_info.append(personal_info.website)
        
        if contact_info:
            contact_para = doc.add_paragraph()
            contact_para.style = 'Synergie Body'
            contact_run = contact_para.add_run(' | '.join(contact_info))
            contact_run.font.name = self.font_family
            contact_run.font.size = Pt(self.font_sizes['body'])
            contact_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add spacing after header
        doc.add_paragraph()
    
    def add_section_header(self, doc: Document, section_name: str) -> None:
        """
        Add section header to document
        
        Args:
            doc: Document object
            section_name: Name of the section
        """
        
        # Add section header
        section_para = doc.add_paragraph()
        section_para.style = 'Synergie Section'
        section_run = section_para.add_run(section_name)
        section_run.font.name = self.font_family
        section_run.font.size = Pt(self.font_sizes['section'])
        section_run.font.bold = True
        section_run.font.color.rgb = RGBColor(*self.synergie_orange)
        
        # Add spacing after section header
        doc.add_paragraph()
    
    def add_subsection_header(self, doc: Document, subsection_name: str) -> None:
        """
        Add subsection header to document
        
        Args:
            doc: Document object
            subsection_name: Name of the subsection
        """
        
        # Add subsection header
        subsection_para = doc.add_paragraph()
        subsection_para.style = 'Synergie Subsection'
        subsection_run = subsection_para.add_run(subsection_name)
        subsection_run.font.name = self.font_family
        subsection_run.font.size = Pt(self.font_sizes['subsection'])
        subsection_run.font.bold = True
        subsection_run.font.color.rgb = RGBColor(*self.text_color)
    
    def add_body_text(self, doc: Document, text: str) -> None:
        """
        Add body text to document
        
        Args:
            doc: Document object
            text: Text to add
        """
        
        if not text or not text.strip():
            return
        
        # Add body text
        body_para = doc.add_paragraph()
        body_para.style = 'Synergie Body'
        body_run = body_para.add_run(text.strip())
        body_run.font.name = self.font_family
        body_run.font.size = Pt(self.font_sizes['body'])
        body_run.font.color.rgb = RGBColor(*self.text_color)
    
    def add_bullet_point(self, doc: Document, text: str) -> None:
        """
        Add bullet point to document
        
        Args:
            doc: Document object
            text: Text for bullet point
        """
        
        if not text or not text.strip():
            return
        
        # Add bullet point
        bullet_para = doc.add_paragraph()
        bullet_para.style = 'Synergie Body'
        bullet_para.paragraph_format.left_indent = Inches(0.25)
        
        bullet_run = bullet_para.add_run(f"• {text.strip()}")
        bullet_run.font.name = self.font_family
        bullet_run.font.size = Pt(self.font_sizes['body'])
        bullet_run.font.color.rgb = RGBColor(*self.text_color)
    
    def save_document(self, doc: Document, file_path: str) -> bool:
        """
        Save document to file
        
        Args:
            doc: Document object
            file_path: Path to save the document
            
        Returns:
            True if successful, False otherwise
        """
        
        try:
            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Save document
            doc.save(file_path)
            
            return True
        
        except Exception as e:
            print(f"Error saving document: {str(e)}")
            return False
    
    def get_document_info(self, doc: Document) -> Dict[str, Any]:
        """
        Get document information
        
        Args:
            doc: Document object
            
        Returns:
            Dictionary with document information
        """
        
        return {
            'page_width': doc.sections[0].page_width,
            'page_height': doc.sections[0].page_height,
            'margin_top': doc.sections[0].top_margin,
            'margin_bottom': doc.sections[0].bottom_margin,
            'margin_left': doc.sections[0].left_margin,
            'margin_right': doc.sections[0].right_margin,
            'paragraph_count': len(doc.paragraphs),
            'styles_count': len(doc.styles)
        }
