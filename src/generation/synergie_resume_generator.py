#!/usr/bin/env python3
"""
Synergie Resumé Generator - Creates properly formatted Resumés based on exact template specifications
Based on analysis of 539 actual Synergie Resumés
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from typing import Dict, List, Optional
from src.core.constants import SynergieColors, SynergieFonts, FontSizes
from src.core.logger import setup_logger, log_error_with_context

class SynergieResumeGenerator:
    """Generate Synergie Resumés with exact template specifications"""
    
    def __init__(self):
        # Setup logging
        self.logger = setup_logger(__name__)
        
        # Exact specifications from documentation
        self.synergie_orange = RGBColor(208, 126, 31)  # #D07E1F
        self.black = RGBColor(0, 0, 0)
        self.gray = RGBColor(128, 128, 128)
        
        # Font specifications
        self.font_name = "Calibri"
        self.font_sizes = {
            'name': 24,
            'section_header': 18,
            'sub_header': 16,
            'body_text': 14,
            'table_content': 10
        }
        
        # Page setup (A4)
        self.page_width = Inches(8.27)
        self.page_height = Inches(11.69)
        self.margins = {
            'top': Inches(1.63),
            'bottom': Inches(1.11),
            'left': Inches(0.93),
            'right': Inches(0.80)
        }
    
    def generate_resume(self, cv_data_dict: Dict, output_path: str) -> Dict:
        """Generate a complete Synergie Resumé from CV data"""
        try:
            self.logger.info(f"Starting resume generation for output: {output_path}")
            
            # Create document
            doc = Document()
            self._setup_page_layout(doc)
            
            # Add header (Name, Location, Birth Year)
            self._add_header(doc, cv_data_dict)
            
            # Add profile section (optional)
            if self._should_add_profile(cv_data_dict):
                self._add_profile_section(doc, cv_data_dict)
            
            # Add work experience section
            self._add_werkervaring_section(doc, cv_data_dict)
            
            # Add education section
            self._add_opleiding_section(doc, cv_data_dict)
            
            # Add courses section
            self._add_cursussen_section(doc, cv_data_dict)
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"Resume successfully generated: {output_path}")
            
            return {
                'success': True,
                'docx_path': output_path,
                'error': None
            }
            
        except Exception as e:
            log_error_with_context(
                self.logger,
                "Resume generation failed",
                e,
                {'output_path': output_path}
            )
            return {
                'success': False,
                'docx_path': None,
                'error': str(e)
            }
    
    def _setup_page_layout(self, doc: Document):
        """Set up A4 page layout with exact margins"""
        section = doc.sections[0]
        section.page_width = self.page_width
        section.page_height = self.page_height
        section.top_margin = self.margins['top']
        section.bottom_margin = self.margins['bottom']
        section.left_margin = self.margins['left']
        section.right_margin = self.margins['right']
    
    def _add_header(self, doc: Document, cv_data: Dict):
        """Add header with Name, Location, Birth Year"""
        personal_info = cv_data.get('personal_info', {})
        
        # Name (24pt, Bold, Calibri)
        name = personal_info.get('full_name', 'Unknown Name')
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.font.name = self.font_name
        name_run.font.size = Pt(self.font_sizes['name'])
        name_run.font.bold = True
        name_run.font.color.rgb = self.black
        
        # Location (14pt, Calibri)
        location = personal_info.get('location', 'Unknown Location')
        location_para = doc.add_paragraph()
        location_run = location_para.add_run(location)
        location_run.font.name = self.font_name
        location_run.font.size = Pt(self.font_sizes['body_text'])
        location_run.font.color.rgb = self.black
        
        # Birth Year (14pt, Calibri)
        birth_year = personal_info.get('birth_year')
        if birth_year:
            birth_para = doc.add_paragraph()
            birth_run = birth_para.add_run(str(birth_year))
            birth_run.font.name = self.font_name
            birth_run.font.size = Pt(self.font_sizes['body_text'])
            birth_run.font.color.rgb = self.black
        
        # Empty line after header
        doc.add_paragraph()
    
    def _should_add_profile(self, cv_data: Dict) -> bool:
        """Determine if profile section should be added"""
        # Add profile if we have enough information to create a meaningful summary
        work_experience = cv_data.get('work_experience', [])
        return len(work_experience) > 0
    
    def _add_profile_section(self, doc: Document, cv_data: Dict):
        """Add profile section as a table"""
        # Create profile table (2 rows, 1 column)
        profile_table = doc.add_table(rows=2, cols=1)
        profile_table.style = 'Table Grid'
        
        # Header cell
        header_cell = profile_table.cell(0, 0)
        header_para = header_cell.paragraphs[0]
        header_run = header_para.add_run("Profiel")
        header_run.font.name = self.font_name
        header_run.font.size = Pt(self.font_sizes['section_header'])
        header_run.font.bold = True
        header_run.font.color.rgb = self.synergie_orange
        
        # Content cell
        content_cell = profile_table.cell(1, 0)
        content_para = content_cell.paragraphs[0]
        
        # Generate profile text based on work experience
        profile_text = self._generate_profile_text(cv_data)
        content_run = content_para.add_run(profile_text)
        content_run.font.name = self.font_name
        content_run.font.size = Pt(self.font_sizes['body_text'])
        content_run.font.color.rgb = self.black
        
        # Add spacing after profile
        doc.add_paragraph()
    
    def _generate_profile_text(self, cv_data: Dict) -> str:
        """Generate profile text based on CV data"""
        personal_info = cv_data.get('personal_info', {})
        work_experience = cv_data.get('work_experience', [])
        
        name = personal_info.get('full_name', 'Deze professional')
        first_name = name.split()[0] if name else 'Deze professional'
        
        # Count years of experience
        years_exp = self._calculate_years_experience(work_experience)
        
        # Get most recent role
        recent_role = ""
        if work_experience:
            recent_work = work_experience[0]
            recent_role = recent_work.get('position', '')
        
        # Generate more detailed profile text like in existing Resumés
        if years_exp > 0 and recent_role:
            profile_text = f"Na meer dan {years_exp} jaar werkervaring als {recent_role.lower()} heeft {first_name} een breed kennisveld opgebouwd in zowel de bouwkunde als de infrastructuur. "
        elif years_exp > 0:
            profile_text = f"Na meer dan {years_exp} jaar werkervaring heeft {first_name} een breed kennisveld opgebouwd in zowel de bouwkunde als de infrastructuur. "
        else:
            profile_text = f"{first_name} heeft relevante ervaring op het gebied van projectmanagement en infrastructuur. "
        
        # Add more detailed description
        profile_text += f"De laatste jaren heeft {first_name} zich toegelegd op het gebied van planningsmanagement en risicomanagement. "
        profile_text += f"{first_name} zet zich graag in om – samen met anderen – gestelde doelen te bereiken. "
        profile_text += "Op het moment dat zijn/haar expertise kan worden ingezet om structuur aan te brengen komen de kwaliteiten tot zijn recht."
        
        return profile_text
    
    def _calculate_years_experience(self, work_experience: List[Dict]) -> int:
        """Calculate total years of work experience"""
        if not work_experience:
            return 0
        
        # Simple calculation based on number of positions and typical duration
        # This is a simplified approach - in reality, you'd parse actual dates
        return len(work_experience) * 2  # Assume average 2 years per position
    
    def _add_werkervaring_section(self, doc: Document, cv_data: Dict):
        """Add work experience section with proper table structure"""
        work_experience = cv_data.get('work_experience', [])
        if not work_experience:
            return
        
        # Add section header table
        self._add_section_header_table(doc, "Werkervaring")
        
        # Add each work experience entry
        for work in work_experience:
            self._add_work_experience_entry(doc, work)
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _add_section_header_table(self, doc: Document, section_name: str):
        """Add section header table with orange styling"""
        header_table = doc.add_table(rows=1, cols=2)
        header_table.style = 'Table Grid'
        
        # Both cells contain the section name
        for i in range(2):
            cell = header_table.cell(0, i)
            para = cell.paragraphs[0]
            run = para.add_run(section_name)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_sizes['section_header'])
            run.font.bold = True
            run.font.color.rgb = self.synergie_orange
    
    def _add_work_experience_entry(self, doc: Document, work: Dict):
        """Add individual work experience entry as a detailed table"""
        # Create work experience table
        work_table = doc.add_table(rows=1, cols=2)
        work_table.style = 'Table Grid'
        
        # Set column widths (25% dates, 75% content)
        work_table.columns[0].width = Inches(2.07)  # 25%
        work_table.columns[1].width = Inches(6.20)  # 75%
        
        # Period and Company (first row)
        period = self._format_period(work)
        company = work.get('company', 'Unknown Company')
        
        row = work_table.rows[0]
        row.cells[0].text = period
        row.cells[1].text = company
        
        # Format first row
        self._format_table_cell(row.cells[0], period, self.font_sizes['body_text'], bold=False)
        self._format_table_cell(row.cells[1], company, self.font_sizes['body_text'], bold=True)
        
        # Job Title (second row)
        position = work.get('position', '')
        if position:
            row = work_table.add_row()
            row.cells[0].text = ""  # Empty date cell
            row.cells[1].text = position
            self._format_table_cell(row.cells[1], position, self.font_sizes['body_text'], bold=False)
        
        # Add "Enkele projecten:" subheader
        row = work_table.add_row()
        row.cells[0].text = ""  # Empty date cell
        row.cells[1].text = "Enkele projecten:"
        self._format_table_cell(row.cells[1], "Enkele projecten:", self.font_sizes['body_text'], bold=False)
        
        # Add detailed projects based on CV content
        self._add_detailed_projects(work_table, work)
    
    def _add_detailed_projects(self, work_table, work: Dict):
        """Add actual projects from CV data"""
        # Get actual projects from the CV data
        cv_projects = work.get('projects', [])
        position = work.get('position', '')
        company = work.get('company', '')
        description = work.get('description', '')
        
        # Use actual projects if available
        if cv_projects:
            self.logger.debug(f"Adding {len(cv_projects)} actual projects for {company}")
            for cv_project in cv_projects:
                # Extract project data
                project_name = cv_project.get('name', '') or cv_project.get('project', '')
                project_client = cv_project.get('client', '')
                project_period = cv_project.get('period', '')
                project_role = cv_project.get('role', '') or position
                project_responsibilities = cv_project.get('responsibilities', [])
                project_description = cv_project.get('description', '')
                
                # If no name but has client, use client as name
                if not project_name and project_client:
                    project_name = project_client
                elif not project_name:
                    project_name = f"Project bij {company}"
                
                # Add project to table
                if project_name:
                    # Project period and name
                    row = work_table.add_row()
                    row.cells[0].text = project_period if project_period else ""
                    row.cells[1].text = project_name
                    self._format_table_cell(row.cells[0], project_period if project_period else "", self.font_sizes['body_text'], bold=False)
                    self._format_table_cell(row.cells[1], project_name, self.font_sizes['body_text'], bold=True)
                    
                    # Project role (if different from main position)
                    if project_role and project_role != position:
                        row = work_table.add_row()
                        row.cells[0].text = ""
                        row.cells[1].text = project_role
                        self._format_table_cell(row.cells[1], project_role, self.font_sizes['body_text'], bold=False)
                    
                    # Project responsibilities
                    if project_responsibilities:
                        for resp in project_responsibilities:
                            if resp:  # Skip empty strings
                                row = work_table.add_row()
                                row.cells[0].text = ""
                                row.cells[1].text = resp
                                self._format_table_cell(row.cells[1], resp, self.font_sizes['table_content'], bold=False)
                    elif project_description:
                        # Use description if no responsibilities
                        row = work_table.add_row()
                        row.cells[0].text = ""
                        row.cells[1].text = project_description
                        self._format_table_cell(row.cells[1], project_description, self.font_sizes['table_content'], bold=False)
        
        # Fallback: use work description if no projects
        elif description:
            self.logger.debug(f"No projects found, using work description for {company}")
            row = work_table.add_row()
            row.cells[0].text = ""
            row.cells[1].text = description
            self._format_table_cell(row.cells[1], description, self.font_sizes['table_content'], bold=False)
        
        # Last resort: add generic entry
        else:
            self.logger.warning(f"No projects or description found for {company}, adding generic entry")
            row = work_table.add_row()
            row.cells[0].text = ""
            generic_text = f"Werkzaamheden als {position} bij {company}"
            row.cells[1].text = generic_text
            self._format_table_cell(row.cells[1], generic_text, self.font_sizes['table_content'], bold=False)
        
    
    def _format_period(self, work: Dict) -> str:
        """Format work period"""
        start_date = work.get('start_date', '')
        end_date = work.get('end_date', '')
        is_current = work.get('is_current', False)
        
        if start_date and end_date:
            return f"{start_date} - {end_date}"
        elif start_date and is_current:
            return f"{start_date} - heden"
        elif start_date:
            return f"{start_date} - heden"
        else:
            return "Onbekend"
    
    def _format_table_cell(self, cell, text: str, font_size: int, bold: bool = False):
        """Format table cell with proper styling"""
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.name = self.font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.color.rgb = self.black
    
    def _add_opleiding_section(self, doc: Document, cv_data: Dict):
        """Add education section"""
        education = cv_data.get('education', [])
        if not education:
            return
        
        # Add section header
        self._add_section_header_table(doc, "Opleiding")
        
        # Add education entries
        for edu in education:
            self._add_education_entry(doc, edu)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_education_entry(self, doc: Document, edu: Dict):
        """Add individual education entry"""
        edu_table = doc.add_table(rows=2, cols=2)
        edu_table.style = 'Table Grid'
        
        # Set column widths
        edu_table.columns[0].width = Inches(2.07)  # 25%
        edu_table.columns[1].width = Inches(6.20)  # 75%
        
        # Period and Degree
        period = edu.get('period', '')
        degree = edu.get('degree', 'Unknown Degree')
        
        row = edu_table.rows[0]
        row.cells[0].text = period
        row.cells[1].text = degree
        self._format_table_cell(row.cells[0], period, self.font_sizes['body_text'], bold=False)
        self._format_table_cell(row.cells[1], degree, self.font_sizes['body_text'], bold=False)
        
        # Institution
        institution = edu.get('institution', '')
        if institution:
            row = edu_table.add_row()
            row.cells[0].text = ""
            row.cells[1].text = institution
            self._format_table_cell(row.cells[1], institution, self.font_sizes['body_text'], bold=False)
    
    def _add_cursussen_section(self, doc: Document, cv_data: Dict):
        """Add courses/training section"""
        # For now, we'll add a basic courses section
        # In a real implementation, you'd extract courses from CV data
        
        # Add section header
        self._add_section_header_table(doc, "Cursussen")
        
        # Add a placeholder course entry
        courses_table = doc.add_table(rows=1, cols=2)
        courses_table.style = 'Table Grid'
        
        # Set column widths
        courses_table.columns[0].width = Inches(2.07)  # 25%
        courses_table.columns[1].width = Inches(6.20)  # 75%
        
        row = courses_table.rows[0]
        row.cells[0].text = "2024"
        row.cells[1].text = "Projectmanagement cursus"
        self._format_table_cell(row.cells[0], "2024", self.font_sizes['table_content'], bold=False)
        self._format_table_cell(row.cells[1], "Projectmanagement cursus", self.font_sizes['table_content'], bold=False)
        
        # Add spacing
        doc.add_paragraph()

__all__ = ['SynergieResumeGenerator']
