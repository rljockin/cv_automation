#!/usr/bin/env python3
"""
Table Builder - Build 2-column tables for Synergie Resumé layout
Creates complex table structures with exact specifications
"""

from typing import List, Optional, Tuple, Dict, Any
from datetime import datetime, date

from docx import Document
from docx.table import Table, _Cell
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.core import (
    SynergieTableSpecs,
    SynergieFonts,
    SynergieColors,
    WorkExperience,
    PersonalInfo,
    Language
)


class TableBuilder:
    """
    Build 2-column tables for Synergie Resumé layout
    
    COMPLEX COMPONENT - handles:
    - 2-column table structure (0.89" dates, 5.88" content)
    - Work experience tables
    - Education tables
    - Skills tables
    - Project tables
    - Exact cell formatting and alignment
    
    Based on comprehensive Resumé analysis:
    - 539 Resumés analyzed
    - Exact table specifications extracted
    - Pixel-perfect matching required
    """
    
    def __init__(self):
        """Initialize table builder with Synergie specifications"""
        
        # Table specifications from analysis
        self.date_column_width = SynergieTableSpecs.DATE_COLUMN_WIDTH
        self.content_column_width = SynergieTableSpecs.CONTENT_COLUMN_WIDTH
        self.table_alignment = SynergieTableSpecs.TABLE_ALIGNMENT
        
        # Font specifications
        self.font_family = SynergieFonts.FONT_FAMILY
        self.font_sizes = SynergieFonts.FONT_SIZES
        
        # Color specifications
        self.synergie_orange = SynergieColors.SYNERGIE_ORANGE
        self.text_color = SynergieColors.TEXT_COLOR
    
    def create_work_experience_table(self, doc: Document, work_experiences: List[WorkExperience]) -> Table:
        """
        Create work experience table
        
        Args:
            doc: Document object
            work_experiences: List of work experience objects
            
        Returns:
            Table object with work experience data
        """
        
        if not work_experiences:
            return None
        
        # Create table with 2 columns
        table = doc.add_table(rows=len(work_experiences), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Set column widths
        self._set_table_column_widths(table)
        
        # Add work experience data
        for i, work_exp in enumerate(work_experiences):
            self._add_work_experience_row(table, i, work_exp)
        
        # Apply table styling
        self._apply_table_styling(table)
        
        return table
    
    def create_education_table(self, doc: Document, education_data: List[Dict]) -> Table:
        """
        Create education table
        
        Args:
            doc: Document object
            education_data: List of education entries
            
        Returns:
            Table object with education data
        """
        
        if not education_data:
            return None
        
        # Create table with 2 columns
        table = doc.add_table(rows=len(education_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Set column widths
        self._set_table_column_widths(table)
        
        # Add education data
        for i, edu in enumerate(education_data):
            self._add_education_row(table, i, edu)
        
        # Apply table styling
        self._apply_table_styling(table)
        
        return table
    
    def create_skills_table(self, doc: Document, skills_data: Dict[str, List[str]]) -> Table:
        """
        Create skills table
        
        Args:
            doc: Document object
            skills_data: Dictionary with skill categories and items
            
        Returns:
            Table object with skills data
        """
        
        if not skills_data:
            return None
        
        # Calculate total rows needed
        total_rows = sum(len(items) for items in skills_data.values())
        
        if total_rows == 0:
            return None
        
        # Create table with 2 columns
        table = doc.add_table(rows=total_rows, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Set column widths
        self._set_table_column_widths(table)
        
        # Add skills data
        row_index = 0
        for category, items in skills_data.items():
            for item in items:
                self._add_skill_row(table, row_index, category, item)
                row_index += 1
        
        # Apply table styling
        self._apply_table_styling(table)
        
        return table
    
    def create_projects_table(self, doc: Document, projects_data: List[Dict]) -> Table:
        """
        Create projects table
        
        Args:
            doc: Document object
            projects_data: List of project entries
            
        Returns:
            Table object with projects data
        """
        
        if not projects_data:
            return None
        
        # Create table with 2 columns
        table = doc.add_table(rows=len(projects_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Set column widths
        self._set_table_column_widths(table)
        
        # Add projects data
        for i, project in enumerate(projects_data):
            self._add_project_row(table, i, project)
        
        # Apply table styling
        self._apply_table_styling(table)
        
        return table
    
    def _set_table_column_widths(self, table: Table) -> None:
        """Set exact column widths for Synergie table"""
        
        # Set column widths
        table.columns[0].width = Inches(self.date_column_width)
        table.columns[1].width = Inches(self.content_column_width)
    
    def _add_work_experience_row(self, table: Table, row_index: int, work_exp: WorkExperience) -> None:
        """Add work experience row to table"""
        
        # Get cells
        date_cell = table.cell(row_index, 0)
        content_cell = table.cell(row_index, 1)
        
        # Format date cell
        self._format_date_cell(date_cell, work_exp)
        
        # Format content cell
        self._format_work_experience_content_cell(content_cell, work_exp)
    
    def _format_date_cell(self, cell: _Cell, work_exp: WorkExperience) -> None:
        """Format date cell with work experience dates"""
        
        # Clear existing content
        cell.text = ''
        
        # Format dates
        date_text = self._format_work_experience_dates(work_exp)
        
        # Add date paragraph
        date_para = cell.paragraphs[0]
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_para.add_run(date_text)
        date_run.font.name = self.font_family
        date_run.font.size = Pt(self.font_sizes['body'])
        date_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Set cell properties
        self._set_cell_properties(cell)
    
    def _format_work_experience_dates(self, work_exp: WorkExperience) -> str:
        """Format work experience dates"""
        
        if not work_exp.start_date:
            return ""
        
        start_year = work_exp.start_date.year
        
        if work_exp.is_current:
            return f"{start_year} - heden"
        elif work_exp.end_date:
            end_year = work_exp.end_date.year
            return f"{start_year} - {end_year}"
        else:
            return str(start_year)
    
    def _format_work_experience_content_cell(self, cell: _Cell, work_exp: WorkExperience) -> None:
        """Format work experience content cell"""
        
        # Clear existing content
        cell.text = ''
        
        # Add company and position
        if work_exp.company and work_exp.position:
            company_para = cell.paragraphs[0]
            company_run = company_para.add_run(f"{work_exp.position} bij {work_exp.company}")
            company_run.font.name = self.font_family
            company_run.font.size = Pt(self.font_sizes['subsection'])
            company_run.font.bold = True
            company_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add location if available
        if work_exp.location:
            location_para = cell.add_paragraph()
            location_run = location_para.add_run(work_exp.location)
            location_run.font.name = self.font_family
            location_run.font.size = Pt(self.font_sizes['body'])
            location_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add description if available
        if work_exp.description:
            desc_para = cell.add_paragraph()
            desc_run = desc_para.add_run(work_exp.description)
            desc_run.font.name = self.font_family
            desc_run.font.size = Pt(self.font_sizes['body'])
            desc_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add responsibilities
        if work_exp.responsibilities:
            for responsibility in work_exp.responsibilities[:3]:  # Limit to 3
                resp_para = cell.add_paragraph()
                resp_para.paragraph_format.left_indent = Inches(0.1)
                resp_run = resp_para.add_run(f"• {responsibility}")
                resp_run.font.name = self.font_family
                resp_run.font.size = Pt(self.font_sizes['body'])
                resp_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Set cell properties
        self._set_cell_properties(cell)
    
    def _add_education_row(self, table: Table, row_index: int, education: Dict) -> None:
        """Add education row to table"""
        
        # Get cells
        date_cell = table.cell(row_index, 0)
        content_cell = table.cell(row_index, 1)
        
        # Format date cell
        self._format_education_date_cell(date_cell, education)
        
        # Format content cell
        self._format_education_content_cell(content_cell, education)
    
    def _format_education_date_cell(self, cell: _Cell, education: Dict) -> None:
        """Format education date cell"""
        
        # Clear existing content
        cell.text = ''
        
        # Format dates
        date_text = self._format_education_dates(education)
        
        # Add date paragraph
        date_para = cell.paragraphs[0]
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_para.add_run(date_text)
        date_run.font.name = self.font_family
        date_run.font.size = Pt(self.font_sizes['body'])
        date_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Set cell properties
        self._set_cell_properties(cell)
    
    def _format_education_dates(self, education: Dict) -> str:
        """Format education dates"""
        
        if 'start_year' in education and 'end_year' in education:
            return f"{education['start_year']} - {education['end_year']}"
        elif 'year' in education:
            return str(education['year'])
        else:
            return ""
    
    def _format_education_content_cell(self, cell: _Cell, education: Dict) -> None:
        """Format education content cell"""
        
        # Clear existing content
        cell.text = ''
        
        # Add degree and institution
        if 'degree' in education and 'institution' in education:
            degree_para = cell.paragraphs[0]
            degree_run = degree_para.add_run(f"{education['degree']} - {education['institution']}")
            degree_run.font.name = self.font_family
            degree_run.font.size = Pt(self.font_sizes['subsection'])
            degree_run.font.bold = True
            degree_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add specialization if available
        if 'specialization' in education:
            spec_para = cell.add_paragraph()
            spec_run = spec_para.add_run(education['specialization'])
            spec_run.font.name = self.font_family
            spec_run.font.size = Pt(self.font_sizes['body'])
            spec_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Set cell properties
        self._set_cell_properties(cell)
    
    def _add_skill_row(self, table: Table, row_index: int, category: str, item: str) -> None:
        """Add skill row to table"""
        
        # Get cells
        date_cell = table.cell(row_index, 0)
        content_cell = table.cell(row_index, 1)
        
        # Format date cell (empty for skills)
        self._format_empty_date_cell(date_cell)
        
        # Format content cell
        self._format_skill_content_cell(content_cell, category, item)
    
    def _format_empty_date_cell(self, cell: _Cell) -> None:
        """Format empty date cell"""
        
        cell.text = ''
        self._set_cell_properties(cell)
    
    def _format_skill_content_cell(self, cell: _Cell, category: str, item: str) -> None:
        """Format skill content cell"""
        
        # Clear existing content
        cell.text = ''
        
        # Add skill item
        skill_para = cell.paragraphs[0]
        skill_run = skill_para.add_run(item)
        skill_run.font.name = self.font_family
        skill_run.font.size = Pt(self.font_sizes['body'])
        skill_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Set cell properties
        self._set_cell_properties(cell)
    
    def _add_project_row(self, table: Table, row_index: int, project: Dict) -> None:
        """Add project row to table"""
        
        # Get cells
        date_cell = table.cell(row_index, 0)
        content_cell = table.cell(row_index, 1)
        
        # Format date cell
        self._format_project_date_cell(date_cell, project)
        
        # Format content cell
        self._format_project_content_cell(content_cell, project)
    
    def _format_project_date_cell(self, cell: _Cell, project: Dict) -> None:
        """Format project date cell"""
        
        # Clear existing content
        cell.text = ''
        
        # Format dates
        date_text = self._format_project_dates(project)
        
        # Add date paragraph
        date_para = cell.paragraphs[0]
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_para.add_run(date_text)
        date_run.font.name = self.font_family
        date_run.font.size = Pt(self.font_sizes['body'])
        date_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Set cell properties
        self._set_cell_properties(cell)
    
    def _format_project_dates(self, project: Dict) -> str:
        """Format project dates"""
        
        if 'start_year' in project and 'end_year' in project:
            return f"{project['start_year']} - {project['end_year']}"
        elif 'year' in project:
            return str(project['year'])
        else:
            return ""
    
    def _format_project_content_cell(self, cell: _Cell, project: Dict) -> None:
        """Format project content cell"""
        
        # Clear existing content
        cell.text = ''
        
        # Add project name
        if 'name' in project:
            name_para = cell.paragraphs[0]
            name_run = name_para.add_run(project['name'])
            name_run.font.name = self.font_family
            name_run.font.size = Pt(self.font_sizes['subsection'])
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add project description
        if 'description' in project:
            desc_para = cell.add_paragraph()
            desc_run = desc_para.add_run(project['description'])
            desc_run.font.name = self.font_family
            desc_run.font.size = Pt(self.font_sizes['body'])
            desc_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Add client if available
        if 'client' in project:
            client_para = cell.add_paragraph()
            client_run = client_para.add_run(f"Klant: {project['client']}")
            client_run.font.name = self.font_family
            client_run.font.size = Pt(self.font_sizes['body'])
            client_run.font.color.rgb = RGBColor(*self.text_color)
        
        # Set cell properties
        self._set_cell_properties(cell)
    
    def _set_cell_properties(self, cell: _Cell) -> None:
        """Set cell properties for consistent formatting"""
        
        # Set vertical alignment
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        
        # Remove cell borders (Synergie style)
        self._remove_cell_borders(cell)
    
    def _remove_cell_borders(self, cell: _Cell) -> None:
        """Remove cell borders for clean Synergie look"""
        
        # Get cell properties
        tc = cell._tc
        tcPr = tc.tcPr
        
        # Remove borders
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
        
        # Set all borders to none
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
    
    def _apply_table_styling(self, table: Table) -> None:
        """Apply overall table styling"""
        
        # Set table alignment
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Remove table borders
        self._remove_table_borders(table)
    
    def _remove_table_borders(self, table: Table) -> None:
        """Remove table borders for clean Synergie look"""
        
        # Get table properties
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Remove borders
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
        
        # Set all borders to none
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
