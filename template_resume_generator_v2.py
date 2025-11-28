#!/usr/bin/env python3
"""
Template-Based Resumé Generator V2
Uses the Synergie template with proper table structure
"""

import os
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.core.utils import capitalize_name, extract_name_components

class TemplateResumeGeneratorV2:
    """Generate Resumés using Synergie template with proper table structure"""
    
    def __init__(self, template_path: str):
        self.template_path = template_path
        self._name_prefixes = {
            'van', 'de', 'der', 'den', 'het', 'ter', 'te', 'ten', 'op',
            'aan', 'bij', 'voor', "d'", 'la', 'le', "van't", "v.d.", "v/d"
        }
        
        # Synergie brand colors (updated October 2025)
        self.synergie_orange = RGBColor(208, 126, 31)  # #D07E1F - Name, Location, Birthyear, Opdrachtgever, Project
        self.teal = RGBColor(0, 146, 159)  # #00929F - Section titles, Functie, Bullets
        self.black = RGBColor(0, 0, 0)  # #000000 - All other text
        self.gray = RGBColor(128, 128, 128)  # #808080 - secondary text
        
        # Font specifications (updated October 2025)
        self.font_name = "Corbel"  # Primary font for all Resumés
        self.font_sizes = {
            'name': 18,           # Name (18pt)
            'location': 14,       # Location (14pt)
            'birth_year': 14,     # Birth year (14pt)
            'section_header': 10, # Section headers (10pt)
            'body_text': 10       # All other text (10pt)
        }
        
        # Page setup (A4 standard from analysis)
        self.page_width = 8.27  # inches
        self.page_height = 11.69  # inches
        self.margins = {
            'top': 1.63,    # inches
            'bottom': 1.11, # inches
            'left': 0.93,   # inches
            'right': 0.80   # inches
        }
        
    
    def generate_resume(self, cv_data: Dict, output_path: str) -> Dict:
        """Generate Resumé using template with proper table structure"""
        try:
            print(f"Generating Resumé using template: {self.template_path}")
            
            # Check if template exists
            if not os.path.exists(self.template_path):
                return {
                    'success': False,
                    'error': f'Template file not found: {self.template_path}'
                }
            
            # Load template
            doc = Document(self.template_path)
            print(f"Template loaded successfully")
            
            # Find content insertion point
            content_area = self._find_content_area(doc)
            if not content_area:
                return {
                    'success': False,
                    'error': 'Could not find content insertion point in template'
                }
            
            # Clear existing content
            self._clear_content_area(content_area)
            
            # Insert CV data with proper table structure
            self._insert_cv_data_with_tables(doc, cv_data, content_area)
            
            # Setup footer for pages 2+ (first page footer remains unchanged)
            self._setup_footer(doc, cv_data)
            
            # Apply Synergie styling
            self._apply_synergie_styling(doc)
            
            # Save document
            doc.save(output_path)
            print(f"Resumé saved to: {output_path}")
            
            return {
                'success': True,
                'output_path': output_path,
                'file_size': os.path.getsize(output_path)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def _find_content_area(self, doc: Document):
        """Find where to insert CV content in the template"""
        # Strategy 1: Look for placeholder text
        placeholders = [
            "CONTENT", "INSERT", "RESUME", "CV", "DATA", 
            "PLACEHOLDER", "TEMPLATE", "HERE", "CONTENT_AREA",
            "NAME", "NAAM", "LOCATION", "WOONPLAATS"
        ]
        
        for para in doc.paragraphs:
            para_text = para.text.strip().upper()
            for placeholder in placeholders:
                if placeholder in para_text:
                    print(f"Found content area marker: '{placeholder}' in paragraph")
                    return para
        
        # Strategy 2: Look for empty paragraphs in the middle of the document
        total_paras = len(doc.paragraphs)
        if total_paras > 4:  # Must have header, content area, and footer
            # Look for empty paragraph in the middle third
            start_idx = total_paras // 3
            end_idx = 2 * total_paras // 3
            
            for i in range(start_idx, end_idx):
                para = doc.paragraphs[i]
                if not para.text.strip():
                    print(f"Found empty paragraph at index {i} for content insertion")
                    return para
        
        # Strategy 3: Use the first paragraph after potential header
        for i in range(2, min(5, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            if not para.text.strip() or len(para.text.strip()) < 50:
                print(f"Using paragraph at index {i} for content insertion")
                return para
        
        # Strategy 4: Create new paragraph at the end (before footer)
        if len(doc.paragraphs) > 0:
            # Insert before the last paragraph (assuming it's footer)
            last_para = doc.paragraphs[-1]
            new_para = last_para._element.getparent().insert_before(
                last_para._element, 
                last_para._element.__class__()
            )
            print("Created new paragraph for content insertion")
            return doc.paragraphs[-2]  # Return the newly created paragraph
        
        return None
    
    def _clear_content_area(self, content_para):
        """Clear existing content from the content area"""
        if content_para:
            content_para.clear()
            print("Content area cleared")
    
    def _insert_cv_data_with_tables(self, doc: Document, cv_data: Dict, content_para):
        """Insert CV data using proper table structure"""
        try:
            personal_info = cv_data.get('personal_info', {})
            work_experience = cv_data.get('work_experience', [])
            education = cv_data.get('education', [])
            courses = cv_data.get('courses', [])
            skills = cv_data.get('skills', [])
            
            # Add name, location, birth year as paragraphs
            name = self._format_header_name(personal_info)
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(name)
            name_run.font.name = self.font_name
            name_run.font.size = Pt(self.font_sizes['name'])
            name_run.font.bold = True
            name_run.font.color.rgb = self.synergie_orange  # #ff8500
            
            # Add location
            location = personal_info.get('location', '')
            if location:
                location_para = doc.add_paragraph()
                location_run = location_para.add_run(location)
                location_run.font.name = self.font_name
                location_run.font.size = Pt(self.font_sizes['location'])
                location_run.font.color.rgb = self.synergie_orange  # #ff8500
            
            # Add birth year
            birth_year = personal_info.get('birth_year')
            if birth_year:
                birth_para = doc.add_paragraph()
                birth_run = birth_para.add_run(str(birth_year))
                birth_run.font.name = self.font_name
                birth_run.font.size = Pt(self.font_sizes['birth_year'])
                birth_run.font.color.rgb = self.synergie_orange  # #ff8500
            
            # Add empty paragraph
            doc.add_paragraph()
            
            # Add profile section (optional)
            profile_summary = cv_data.get('profile_summary', '')
            if profile_summary:
                self._add_profile_table(doc, profile_summary)
            
            # Add work experience section
            if work_experience:
                self._add_werkervaring_tables(doc, work_experience)
            
            # Add education section
            if education:
                self._add_opleiding_tables(doc, education)
            
            # Add courses section (NEW - directly after education)
            if courses:
                self._add_cursussen_tables(doc, courses)
            
            # Add skills section
            if skills:
                self._add_skills_table(doc, skills)
            
            print("CV data inserted successfully with proper table structure")
            
        except Exception as e:
            print(f"Error inserting CV data: {e}")
            raise
    
    def _format_header_name(self, personal_info: Dict) -> str:
        """Return display name with only first and last name (no prefixes/abbreviations)."""
        full_name = personal_info.get('full_name', '') or ''
        first = personal_info.get('first_name')
        last = personal_info.get('last_name')
        
        if not first or not last:
            parsed_first, parsed_last = extract_name_components(full_name)
            first = first or parsed_first
            last = last or parsed_last
        
        display_first = self._normalize_first_name(first)
        display_last = self._normalize_last_name(last)
        
        if not display_first and full_name:
            name_parts = full_name.split() if full_name else []
            if name_parts:
                display_first = self._normalize_first_name(name_parts[0])
        if not display_last and full_name:
            name_parts = full_name.split() if full_name else []
            if name_parts:
                display_last = self._normalize_last_name(name_parts[-1])
        
        display_name = " ".join(part for part in [display_first, display_last] if part)
        return display_name or (full_name or "Unknown Name")
    
    def _normalize_first_name(self, value: Optional[str]) -> str:
        """Use only the first given name and strip abbreviations."""
        if not value:
            return ""
        
        cleaned = value.strip().replace('.', ' ')
        first_token = cleaned.split()[0] if cleaned.split() else ""
        if not first_token:
            return ""
        
        return capitalize_name(first_token)
    
    def _normalize_last_name(self, value: Optional[str]) -> str:
        """Remove Dutch prefixes and abbreviations from last name."""
        if not value:
            return ""
        
        tokens = [token.strip() for token in value.replace('.', ' ').split() if token.strip()]
        filtered = [
            token for token in tokens
            if token.lower() not in self._name_prefixes and len(token) > 1
        ]
        
        if not filtered and tokens:
            filtered = tokens[-1:]
        
        def capitalize_hyphenated(segment: str) -> str:
            return '-'.join(part.capitalize() for part in segment.split('-'))
        
        return " ".join(capitalize_hyphenated(token) for token in filtered)
    
    def format_resume_filename(self, personal_info: Dict) -> str:
        """Generate resume filename in format: Resumé_Synergie projectmanagement_voornaam (tussenvoegsel) achternaam.docx"""
        full_name = personal_info.get('full_name', '') or ''
        
        if not full_name:
            return "Resumé_Synergie projectmanagement_Unknown.docx"
        
        # Split name into parts
        parts = [p.strip() for p in full_name.split() if p.strip()]
        
        if not parts:
            return "Resumé_Synergie projectmanagement_Unknown.docx"
        
        # First name is always the first part
        first_name = parts[0]
        
        if len(parts) == 1:
            # Only first name
            return f"Resumé_Synergie projectmanagement_{first_name}.docx"
        
        # Find prefixes (tussenvoegsels) and last name
        prefixes = []
        last_name_parts = []
        
        i = 1
        while i < len(parts):
            part_lower = parts[i].lower().replace('.', '').replace("'", '')
            # Check if this part is a prefix
            is_prefix = part_lower in self._name_prefixes
            
            if is_prefix:
                prefixes.append(parts[i])
            else:
                # This is the start of the last name
                last_name_parts = parts[i:]
                break
            i += 1
        
        # If we only found prefixes but no last name parts, use the last prefix as last name
        if not last_name_parts and prefixes:
            # Take the last prefix as part of the last name
            last_name_parts = [prefixes[-1]]
            prefixes = prefixes[:-1]
        
        # If still no last name parts, use everything after first name
        if not last_name_parts and len(parts) > 1:
            last_name_parts = parts[1:]
        
        last_name = ' '.join(last_name_parts) if last_name_parts else first_name
        
        # Build filename
        if prefixes:
            prefix_str = ' '.join(prefixes)
            filename = f"Resumé_Synergie projectmanagement_{first_name} ({prefix_str}) {last_name}.docx"
        else:
            filename = f"Resumé_Synergie projectmanagement_{first_name} {last_name}.docx"
        
        # Replace problematic characters for filename
        filename = filename.replace('/', '_').replace('\\', '_').replace(':', '_')
        
        return filename
    
    def _add_profile_table(self, doc: Document, profile_text: str):
        """Add profile section as a simple paragraph (no table)"""
        # Add profile header
        profile_header = doc.add_paragraph()
        profile_header_run = profile_header.add_run("Profielschets")
        profile_header_run.font.name = self.font_name
        profile_header_run.font.size = Pt(self.font_sizes['section_header'])
        profile_header_run.font.bold = True
        profile_header_run.font.color.rgb = self.teal  # #00a6b2
        
        # Add profile content
        profile_content = doc.add_paragraph()
        profile_content_run = profile_content.add_run(profile_text)
        profile_content_run.font.name = self.font_name
        profile_content_run.font.size = Pt(self.font_sizes['body_text'])
        profile_content_run.font.color.rgb = self.black  # #000000
        
        # Add empty paragraph after table
        doc.add_paragraph()
    
    def _prepare_work_experience_groups(self, work_experience: List[Dict]) -> List[Dict]:
        """Prepare consecutively grouped work experiences by employer.
        
        Only groups entries that are directly consecutive (next to each other)
        with the same employer. Non-consecutive entries with the same employer
        are kept separate.
        """
        if not work_experience:
            return []
        
        groups: List[Dict] = []
        
        for idx, entry in enumerate(work_experience):
            raw_company = entry.get('company') or 'Onbekende Werkgever'
            normalized = self._normalize_company_identifier(raw_company)
            display_name = self._clean_company_display(raw_company)
            
            # Check if previous entry is consecutive and same employer
            is_consecutive = False
            if groups and idx > 0:
                prev_entry = work_experience[idx - 1]
                prev_normalized = self._normalize_company_identifier(prev_entry.get('company') or '')
                is_consecutive = (normalized == prev_normalized)
            
            if is_consecutive and groups:
                # Add to existing group
                group = groups[-1]
                group['start_date'] = self._pick_year(group.get('start_date'), entry.get('start_date'), prefer_min=True)
                group['end_date'] = self._pick_year(group.get('end_date'), entry.get('end_date'), prefer_min=False)
                entry_current = entry.get('is_current', False) or not entry.get('end_date')
                group['is_current'] = group.get('is_current', False) or entry_current
            else:
                # Start new group
                group = {
                    'normalized': normalized,
                    'company': display_name,
                    'start_date': entry.get('start_date'),
                    'end_date': entry.get('end_date'),
                    'is_current': entry.get('is_current', False) or not entry.get('end_date'),
                    'items': []
                }
                groups.append(group)
            
            group['items'].extend(self._expand_work_entry(entry))
        
        return groups
    
    def _expand_work_entry(self, work_entry: Dict) -> List[Dict]:
        """Flatten a work entry into renderable project items."""
        items: List[Dict] = []
        projects = work_entry.get('projects') or []
        company_display = self._clean_company_display(work_entry.get('company', ''))
        base_position = work_entry.get('position', '')
        base_responsibilities = work_entry.get('responsibilities') or []
        base_start_date = work_entry.get('start_date')
        base_end_date = work_entry.get('end_date')
        base_is_current = work_entry.get('is_current', False)
        
        if projects:
            for project in projects:
                # Use project period if available, otherwise use work entry period
                project_period = project.get('period')
                if project_period:
                    # Extract dates from period string like "2020-2021" or "2020-heden"
                    period_start, period_end = self._parse_period(project_period)
                    # If parsing failed, fall back to base dates
                    if period_start is None and period_end is None:
                        period_start = base_start_date
                        period_end = base_end_date
                    project_is_current = 'heden' in str(project_period).lower() or 'present' in str(project_period).lower()
                else:
                    period_start = base_start_date
                    period_end = base_end_date
                    project_is_current = base_is_current
                
                responsibilities = project.get('responsibilities') or base_responsibilities
                items.append({
                    'name': project.get('name') or "",
                    'client': project.get('client'),
                    'position': project.get('role') or base_position,
                    'responsibilities': responsibilities or [],
                    'company': company_display,
                    'start_date': period_start,
                    'end_date': period_end,
                    'is_current': project_is_current
                })
        else:
            items.append({
                'name': "",
                'client': None,
                'position': base_position,
                'responsibilities': base_responsibilities or [],
                'company': company_display,
                'start_date': base_start_date,
                'end_date': base_end_date,
                'is_current': base_is_current
            })
        
        return items
    
    def _parse_period(self, period_str: str) -> Tuple[Optional[str], Optional[str]]:
        """Parse period string like '2020-2021' or '2020-heden' into start/end dates."""
        if not period_str:
            return None, None
        
        period_lower = str(period_str).lower().strip()
        
        # Check for "heden" or "present"
        if 'heden' in period_lower or 'present' in period_lower:
            # Extract start year
            match = re.search(r'(19|20)\d{2}', period_lower)
            start_year = match.group(0) if match else None
            return start_year, None
        
        # Try to extract two years
        years = re.findall(r'(19|20)\d{2}', period_lower)
        if len(years) >= 2:
            return years[0], years[1]
        elif len(years) == 1:
            return years[0], years[0]
        
        return None, None
    
    def _normalize_company_identifier(self, company: str) -> str:
        """Normalize company name for grouping comparisons."""
        base = company.split('(')[0]
        return re.sub(r'\s+', ' ', base.strip().lower())
    
    def _clean_company_display(self, company: str) -> str:
        """Return a clean employer name for display."""
        cleaned = company.split('(')[0].strip()
        return cleaned or company or "Unknown Company"
    
    def _pick_year(self, current: Optional[str], candidate: Optional[str], prefer_min: bool) -> Optional[str]:
        """Pick earliest or latest year string between two values."""
        if candidate is None:
            return current
        if current is None:
            return candidate
        
        current_num = self._extract_year_number(current)
        candidate_num = self._extract_year_number(candidate)
        
        if current_num is not None and candidate_num is not None:
            if prefer_min:
                return current if current_num <= candidate_num else candidate
            return current if current_num >= candidate_num else candidate
        
        if current_num is None and candidate_num is not None:
            return candidate if prefer_min else candidate
        if candidate_num is None and current_num is not None:
            return current
        
        return current if prefer_min else candidate
    
    def _extract_year_number(self, value: Optional[str]) -> Optional[int]:
        """Extract numeric year from string value."""
        if not value:
            return None
        match = re.search(r'(19|20)\d{2}', str(value))
        return int(match.group(0)) if match else None
    
    def _format_period_text(self, start_date: Optional[str], end_date: Optional[str], is_current: bool) -> str:
        """Format period text for left column (YYYY or YYYY-YYYY/Heden)."""
        start_year = self._extract_year_number(start_date)
        end_year = self._extract_year_number(end_date)
        
        start_text = str(start_year) if start_year else (start_date[:4] if start_date else "")
        end_text: Optional[str]
        if is_current or not end_date:
            end_text = "Heden"
        else:
            end_text = str(end_year) if end_year else (end_date[:4] if end_date else "")
        
        if start_text and end_text and start_text == end_text:
            return start_text
        
        if start_text and end_text:
            return f"{start_text}-{end_text}"
        
        return start_text or end_text or ""
    
    def _render_responsibilities(self, doc: Document, right_cell, responsibilities: List[str], position: str, company: str) -> None:
        """Render bullet list of responsibilities with teal bullets and black text."""
        from docx.shared import Cm, Pt
        
        cleaned_responsibilities = [resp for resp in responsibilities if resp]
        if cleaned_responsibilities:
            for resp in cleaned_responsibilities:
                detailed_resp = self._expand_responsibility(resp, position, company)
                detailed_resp = self._ensure_period_at_end(detailed_resp)
                resp_para = right_cell.add_paragraph()
                
                # Set up proper bullet list formatting (indent for bullet)
                resp_para.paragraph_format.left_indent = Cm(0.5)
                resp_para.paragraph_format.first_line_indent = Cm(-0.5)
                resp_para.paragraph_format.space_after = Pt(0)
                
                # Add bullet character with teal color (only the bullet, not the text)
                bullet_run = resp_para.add_run("• ")
                bullet_run.font.name = self.font_name
                bullet_run.font.size = Pt(self.font_sizes['body_text'])
                bullet_run.font.color.rgb = self.teal  # Only bullet is teal
                
                # Add the text with black color
                text_run = resp_para.add_run(detailed_resp)
                text_run.font.name = self.font_name
                text_run.font.size = Pt(self.font_sizes['body_text'])
                text_run.font.color.rgb = self.black  # Text is black
        else:
            self._render_default_responsibilities(doc, right_cell, position, company, single=True)
    
    def _render_default_responsibilities(self, doc: Document, right_cell, position: str, company: str, single: bool = False) -> None:
        """Render default responsibilities with teal bullets and black text."""
        from docx.shared import Cm, Pt
        
        position_text = position.lower() if position else "werkzaamheden"
        defaults = [
            f"Uitvoeren van {position_text} taken voor {company}",
            "Verantwoordelijk voor dagelijkse operationele taken",
            "Samenwerken met collega's en stakeholders",
        ]
        
        if single:
            defaults = defaults[:1]
        
        for default_text in defaults:
            default_text = self._ensure_period_at_end(default_text)
            resp_para = right_cell.add_paragraph()
            
            # Set up proper bullet list formatting (indent for bullet)
            resp_para.paragraph_format.left_indent = Cm(0.5)
            resp_para.paragraph_format.first_line_indent = Cm(-0.5)
            resp_para.paragraph_format.space_after = Pt(0)
            
            # Add bullet character with teal color (only the bullet, not the text)
            bullet_run = resp_para.add_run("• ")
            bullet_run.font.name = self.font_name
            bullet_run.font.size = Pt(self.font_sizes['body_text'])
            bullet_run.font.color.rgb = self.teal  # Only bullet is teal
            
            # Add the text with black color
            text_run = resp_para.add_run(default_text)
            text_run.font.name = self.font_name
            text_run.font.size = Pt(self.font_sizes['body_text'])
            text_run.font.color.rgb = self.black  # Text is black
            text_run.font.bold = False
    
    def _add_werkervaring_tables(self, doc: Document, work_experience):
        """Add work experience section using proper table structure"""
        # Add section header table
        header_table = doc.add_table(rows=1, cols=2)
        # Remove table borders - make tables invisible
        header_table.style = None
        
        # Set cell widths individually (this is the key fix!)
        header_table.allow_autofit = False  # Disable autofit to respect our column widths
        header_table.cell(0, 0).width = Inches(1.2)   # Left column (1.2 inch - fits "2024-Heden")
        header_table.cell(0, 1).width = Inches(5.8)   # Wide right column
        
        # Header only in left column
        left_cell = header_table.cell(0, 0)
        para = left_cell.paragraphs[0]
        run = para.add_run("Werkervaring")
        run.font.name = self.font_name
        run.font.size = Pt(self.font_sizes['section_header'])
        run.font.bold = True
        run.font.color.rgb = self.teal  # #00a6b2
        
        # Right column empty
        right_cell = header_table.cell(0, 1)
        right_cell.text = ""
        
        # Add empty paragraph
        doc.add_paragraph()
        
        # Add standard "Zelfstandig Ondernemer" entry as first work experience
        self._render_zelfstandig_ondernemer_entry(doc)
        
        # Add grouped work experience entries
        if work_experience and isinstance(work_experience, list):
            grouped_entries = self._prepare_work_experience_groups(work_experience)
            for group in grouped_entries:
                # Check if this is a bundled group (multiple items) or single entry
                if len(group.get('items', [])) > 1:
                    self._render_bundled_work_experience_group(doc, group)
                else:
                    # Single entry - render normally without bundling
                    self._render_single_work_experience_entry(doc, group.get('items', [])[0] if group.get('items') else {})

    def _render_bundled_work_experience_group(self, doc: Document, group: Dict) -> None:
        """Render a bundled group (one employer, multiple projects) with project years in left column."""
        total_period_text = self._format_period_text(
            group.get('start_date'),
            group.get('end_date'),
            group.get('is_current', False)
        )
        company_display = group.get('company') or 'Onbekende Werkgever'
        items = group.get('items', [])

        # First row: Total period + Company name
        work_table = doc.add_table(rows=1, cols=2)
        work_table.style = None
        work_table.allow_autofit = False
        work_table.cell(0, 0).width = Inches(1.2)
        work_table.cell(0, 1).width = Inches(5.8)

        left_cell = work_table.cell(0, 0)
        left_cell.text = total_period_text
        self._style_table_cell(left_cell, self.font_sizes['body_text'])

        right_cell = work_table.cell(0, 1)
        right_cell.text = ""

        # Company name at top (orange, bold)
        company_para = right_cell.paragraphs[0]
        company_run = company_para.add_run(company_display)
        company_run.font.name = self.font_name
        company_run.font.size = Pt(self.font_sizes['body_text'])
        company_run.font.color.rgb = self.synergie_orange
        company_run.font.bold = True

        # Add each project with its own row (project year in left column, italic)
        for item in items:
            project_name = (item.get('name') or "").strip()
            client = item.get('client')
            display_name = project_name

            if display_name and client and client not in display_name:
                display_name = f"{display_name} ({client})"
            elif not display_name:
                display_name = company_display

            item_start = item.get('start_date')
            item_end = item.get('end_date')
            item_is_current = item.get('is_current', False)
            project_period_text = self._format_period_text(item_start, item_end, item_is_current) if (item_start or item_end) else ""

            # Add new row for this project
            project_row = work_table.add_row()
            project_row.cells[0].width = Inches(1.2)
            project_row.cells[1].width = Inches(5.8)

            # Left column: Project period (italic, standard color, 9pt font size for bundled entries)
            project_left_cell = project_row.cells[0]
            project_left_cell.text = ""
            if project_period_text:
                project_left_para = project_left_cell.paragraphs[0]
                project_left_run = project_left_para.add_run(project_period_text)
                project_left_run.font.name = self.font_name
                project_left_run.font.size = Pt(9)  # 9pt font size for bundled project years (1 size smaller)
                project_left_run.font.color.rgb = self.black  # Standard color
                project_left_run.font.italic = True  # Italic (schuingedrukt)
                project_left_run.font.bold = False
            self._style_table_cell(project_left_cell, 9)  # Use 9pt for cell styling too

            # Right column: Project name, position, responsibilities
            project_right_cell = project_row.cells[1]
            project_right_cell.text = ""

            project_para = project_right_cell.paragraphs[0]
            project_run = project_para.add_run(display_name)
            project_run.font.name = self.font_name
            project_run.font.size = Pt(self.font_sizes['body_text'])
            project_run.font.color.rgb = self.synergie_orange
            project_run.font.bold = True

            position_text = item.get('position', '')
            if position_text:
                position_para = project_right_cell.add_paragraph()
                position_run = position_para.add_run(position_text)
                position_run.font.name = self.font_name
                position_run.font.size = Pt(self.font_sizes['body_text'])
                position_run.font.color.rgb = self.teal
                position_run.font.bold = True

            responsibilities = item.get('responsibilities', [])
            self._render_responsibilities(doc, project_right_cell, responsibilities, position_text, company_display)

        doc.add_paragraph()

    def _render_single_work_experience_entry(self, doc: Document, item: Dict) -> None:
        """Render a single work experience entry (not bundled) in normal format."""
        item_start = item.get('start_date')
        item_end = item.get('end_date')
        item_is_current = item.get('is_current', False)
        period_text = self._format_period_text(item_start, item_end, item_is_current)
        
        company_display = item.get('company') or 'Onbekende Werkgever'
        project_name = (item.get('name') or "").strip()
        client = item.get('client')
        position_text = item.get('position', '')
            
        # Build display name: project name + company (or just company if no project)
        if project_name:
            if client and client not in project_name:
                display_name = f"{project_name} ({client}), {company_display}"
            else:
                display_name = f"{project_name}, {company_display}"
        else:
            display_name = company_display

        work_table = doc.add_table(rows=1, cols=2)
        work_table.style = None
        work_table.allow_autofit = False
        work_table.cell(0, 0).width = Inches(1.2)
        work_table.cell(0, 1).width = Inches(5.8)

        # Left column: Period (normal, not italic)
        left_cell = work_table.cell(0, 0)
        left_cell.text = period_text
        self._style_table_cell(left_cell, self.font_sizes['body_text'])

        # Right column: Project/Company name, position, responsibilities
        right_cell = work_table.cell(0, 1)
        right_cell.text = ""

        first_line_para = right_cell.paragraphs[0]
        first_line_run = first_line_para.add_run(display_name)
        first_line_run.font.name = self.font_name
        first_line_run.font.size = Pt(self.font_sizes['body_text'])
        first_line_run.font.color.rgb = self.synergie_orange
        first_line_run.font.bold = True

        if position_text:
            position_para = right_cell.add_paragraph()
            position_run = position_para.add_run(position_text)
            position_run.font.name = self.font_name
            position_run.font.size = Pt(self.font_sizes['body_text'])
            position_run.font.color.rgb = self.teal
            position_run.font.bold = True

        responsibilities = item.get('responsibilities', [])
        self._render_responsibilities(doc, right_cell, responsibilities, position_text, company_display)

        doc.add_paragraph()
    
    def _render_zelfstandig_ondernemer_entry(self, doc: Document) -> None:
        """Render standard 'Zelfstandig Ondernemer' entry as first work experience."""
        work_table = doc.add_table(rows=1, cols=2)
        work_table.style = None
        work_table.allow_autofit = False
        work_table.cell(0, 0).width = Inches(1.2)
        work_table.cell(0, 1).width = Inches(5.8)

        # Left column: Empty (no period)
        left_cell = work_table.cell(0, 0)
        left_cell.text = ""
        self._style_table_cell(left_cell, self.font_sizes['body_text'])

        # Right column: Company name, position, and "Enkele projecten:"
        right_cell = work_table.cell(0, 1)
        right_cell.text = ""

        # Company name: "Bedrijfsnaam" (orange, bold)
        company_para = right_cell.paragraphs[0]
        company_run = company_para.add_run("Bedrijfsnaam")
        company_run.font.name = self.font_name
        company_run.font.size = Pt(self.font_sizes['body_text'])
        company_run.font.color.rgb = self.synergie_orange
        company_run.font.bold = True

        # Position: "Zelfstandig Ondernemer" (teal/blue, bold)
        position_para = right_cell.add_paragraph()
        position_run = position_para.add_run("Zelfstandig Ondernemer")
        position_run.font.name = self.font_name
        position_run.font.size = Pt(self.font_sizes['body_text'])
        position_run.font.color.rgb = self.teal
        position_run.font.bold = True

        # "Enkele projecten:" (standard black, not bold)
        projects_para = right_cell.add_paragraph()
        projects_run = projects_para.add_run("Enkele projecten:")
        projects_run.font.name = self.font_name
        projects_run.font.size = Pt(self.font_sizes['body_text'])
        projects_run.font.color.rgb = self.black
        projects_run.font.bold = False

        doc.add_paragraph()
    
    def _expand_responsibility(self, resp: str, position: str, company: str) -> str:
        """Return original responsibility text without expansion - use as-is from CV"""
        # Simply return the original responsibility text without any modifications
        return resp
    
    def _ensure_period_at_end(self, text: str) -> str:
        """Ensure text ends with a period if it doesn't already"""
        text = text.strip()
        if not text:
            return text
        # Check if it already ends with punctuation
        if text[-1] not in '.!?':
            return text + '.'
        return text
    
    def _add_opleiding_tables(self, doc: Document, education):
        """Add education section using proper table structure"""
        # Add section header table
        header_table = doc.add_table(rows=1, cols=2)
        # Remove table borders - make tables invisible
        header_table.style = None
        
        # Set cell widths individually (this is the key fix!)
        header_table.allow_autofit = False  # Disable autofit to respect our column widths
        header_table.cell(0, 0).width = Inches(1.2)   # Left column (1.2 inch - fits "2024-Heden")
        header_table.cell(0, 1).width = Inches(5.8)   # Wide right column
        
        # Header only in left column
        left_cell = header_table.cell(0, 0)
        para = left_cell.paragraphs[0]
        run = para.add_run("Opleiding")
        run.font.name = self.font_name
        run.font.size = Pt(self.font_sizes['section_header'])
        run.font.bold = True
        run.font.color.rgb = self.teal  # #00a6b2
        
        # Right column empty
        right_cell = header_table.cell(0, 1)
        right_cell.text = ""
        
        # Add empty paragraph
        doc.add_paragraph()
        
        # Add education entries as a single table
        if education and isinstance(education, list):
            edu_table = doc.add_table(rows=len(education), cols=2)
            # Remove table borders - make tables invisible
            edu_table.style = None
            
            # Set cell widths individually (this is the key fix!)
            edu_table.allow_autofit = False  # Disable autofit to respect our column widths
            
            for i, edu in enumerate(education):
                if not edu:  # Check if edu is None
                    continue
                period = edu.get('period', '')
                degree = edu.get('degree', 'Unknown Degree')
                institution = edu.get('institution', '')
                
                
                # Extract only years from period
                if period and len(period) >= 4:
                    # Extract years from period like "2009 - 2009" or "2008 - 2008"
                    years = period.split(' - ')
                    if len(years) >= 2:
                        start_year = years[0][:4] if len(years[0]) >= 4 else years[0]
                        end_year = years[1][:4] if len(years[1]) >= 4 else years[1]
                        period_text = f"{start_year}-{end_year}" if start_year != end_year else start_year
                    else:
                        period_text = years[0][:4] if len(years[0]) >= 4 else years[0]
                else:
                    period_text = ''
                
                # Combine degree and institution
                content = degree or ''  # Ensure content is never None
                if institution:
                    if content:  # Only add newline if degree exists
                        content += f"\n{institution}"
                    else:
                        content = institution  # Just institution if no degree
                # Set cell widths individually
                edu_table.cell(i, 0).width = Inches(1.2)   # Left column for years (1.2 inch - fits "2024-Heden")
                edu_table.cell(i, 1).width = Inches(5.8)   # Wide right column for content
                
                edu_table.cell(i, 0).text = period_text
                edu_table.cell(i, 1).text = content
                
                # Style the cells with proper colors
                self._style_table_cell(edu_table.cell(i, 0), self.font_sizes['body_text'], self.black)  # Dates - Black
                self._style_table_cell(edu_table.cell(i, 1), self.font_sizes['body_text'], self.black)  # Education text - Black
        
        # Add empty paragraph
        doc.add_paragraph()
    
    def _add_cursussen_tables(self, doc: Document, courses):
        """Add courses section using proper table structure (identical to education)"""
        # Add section header table
        header_table = doc.add_table(rows=1, cols=2)
        # Remove table borders - make tables invisible
        header_table.style = None
        
        # Set cell widths individually (same as education)
        header_table.allow_autofit = False  # Disable autofit to respect our column widths
        header_table.cell(0, 0).width = Inches(1.2)   # Left column (1.2 inch)
        header_table.cell(0, 1).width = Inches(5.8)   # Wide right column
        
        # Header only in left column
        left_cell = header_table.cell(0, 0)
        para = left_cell.paragraphs[0]
        run = para.add_run("Cursussen")
        run.font.name = self.font_name
        run.font.size = Pt(self.font_sizes['section_header'])
        run.font.bold = True
        run.font.color.rgb = self.teal  # Same as education header
        
        # Right column empty
        right_cell = header_table.cell(0, 1)
        right_cell.text = ""
        
        # Add empty paragraph
        doc.add_paragraph()
        
        # Add courses entries as a single table (same structure as education)
        if courses and isinstance(courses, list):
            courses_table = doc.add_table(rows=len(courses), cols=2)
            # Remove table borders - make tables invisible
            courses_table.style = None
            
            # Set cell widths individually
            courses_table.allow_autofit = False  # Disable autofit to respect our column widths
            
            for i, course in enumerate(courses):
                if not course:  # Check if course is None
                    continue
                
                year = course.get('year', '')
                name = course.get('name', 'Onbekende Cursus')
                institution = course.get('institution', '')
                
                # Format year (extract only YYYY if longer string)
                if year and len(year) >= 4:
                    year_text = year[:4]
                else:
                    year_text = year if year else ''
                
                # Combine name and institution (same format as education)
                content = name or ''  # Ensure content is never None
                if institution:
                    if content:  # Only add newline if name exists
                        content += f"\n{institution}"
                    else:
                        content = institution  # Just institution if no name
                
                # Set cell widths individually
                courses_table.cell(i, 0).width = Inches(1.2)   # Left column for year
                courses_table.cell(i, 1).width = Inches(5.8)   # Wide right column for content
                
                courses_table.cell(i, 0).text = year_text
                courses_table.cell(i, 1).text = content
                
                # Style the cells with proper colors (same as education)
                self._style_table_cell(courses_table.cell(i, 0), self.font_sizes['body_text'], self.black)  # Year - Black
                self._style_table_cell(courses_table.cell(i, 1), self.font_sizes['body_text'], self.black)  # Course text - Black
        
        # Add empty paragraph
        doc.add_paragraph()
    
    def _add_skills_table(self, doc: Document, skills):
        """Add skills section as simple paragraphs (no table)"""
        if not skills:
            return
            
        # Add section header as simple paragraph
        skills_header = doc.add_paragraph()
        skills_header_run = skills_header.add_run("Automatiseringskennis")
        skills_header_run.font.name = self.font_name
        skills_header_run.font.size = Pt(self.font_sizes['section_header'])
        skills_header_run.font.bold = True
        skills_header_run.font.color.rgb = self.teal  # #00a6b2
        
        # Add skills content as simple paragraph
        if isinstance(skills, list):
            skills_text = ", ".join(skills[:10])  # Limit to 10 skills
        else:
            skills_text = str(skills) if skills else ""
            
        skills_content = doc.add_paragraph()
        skills_content_run = skills_content.add_run(skills_text)
        skills_content_run.font.name = self.font_name
        skills_content_run.font.size = Pt(self.font_sizes['body_text'])
        skills_content_run.font.color.rgb = self.black  # #000000
        
        # Add empty paragraph
        doc.add_paragraph()
    
    def _style_table_cell(self, cell, font_size: int, color=None, bold: bool = False):
        """Style a table cell with proper font and size"""
        if color is None:
            color = self.black  # Default to black
            
        for para in cell.paragraphs:
            if para.runs:  # Check if there are any runs
                for run in para.runs:
                    run.font.name = self.font_name
                    run.font.size = Pt(font_size)
                    run.font.bold = bold
                    run.font.color.rgb = color
            else:
                # If no runs exist, create one with default text
                run = para.add_run('')
                run.font.name = self.font_name
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run.font.color.rgb = color
    
    def _setup_footer(self, doc: Document, cv_data: Dict):
        """Setup footer for pages 2+ with name, month and year in Dutch"""
        try:
            section = doc.sections[0]
            
            # Ensure different footer for first page (keep first page footer as-is)
            section.different_first_page_header_footer = True
            
            # Get candidate name from CV data
            personal_info = cv_data.get('personal_info', {})
            full_name = personal_info.get('full_name', 'Onbekende Naam')
            
            # Get current month and year in Dutch
            now = datetime.now()
            dutch_months = {
                1: 'januari', 2: 'februari', 3: 'maart', 4: 'april',
                5: 'mei', 6: 'juni', 7: 'juli', 8: 'augustus',
                9: 'september', 10: 'oktober', 11: 'november', 12: 'december'
            }
            month_name = dutch_months[now.month]
            year = now.year
            
            # Format: "Resumé [Naam] - [maand] [jaartal]"
            footer_text = f"Resumé {full_name} - {month_name} {year}"
            
            # Access the regular footer (for pages 2+)
            footer = section.footer
            
            # Footer text color from template analysis: #206177 (dark teal)
            footer_color = RGBColor(32, 97, 119)  # #206177
            
            # The placeholder is in a TABLE, not in paragraphs!
            # Template structure: footer.tables[0].rows[1].cells[0] contains placeholder
            if footer.tables and len(footer.tables) > 0:
                footer_table = footer.tables[0]
                
                # Check if table has row 1 (index 1)
                if len(footer_table.rows) > 1:
                    # Access cell 0 of row 1 (where placeholder text is)
                    placeholder_cell = footer_table.rows[1].cells[0]
                    
                    # Clear all content in this cell
                    for para in placeholder_cell.paragraphs:
                        # Clear all runs
                        for run in list(para.runs):
                            run._element.getparent().remove(run._element)
                    
                    # Get or create first paragraph in this cell
                    if len(placeholder_cell.paragraphs) > 0:
                        para = placeholder_cell.paragraphs[0]
                    else:
                        para = placeholder_cell.add_paragraph()
                    
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add our footer text with correct styling
                    footer_run = para.add_run(footer_text)
                    footer_run.font.name = self.font_name
                    footer_run.font.size = Pt(8)
                    footer_run.font.color.rgb = footer_color  # #206177
                    
                    # Cell 1 (paginanummer) blijft volledig intact - we raken deze niet aan!
                    
                    print(f"Footer set for pages 2+: '{footer_text}' (in table cell)")
                else:
                    print("Warning: Footer table has no row 1, cannot set footer")
            else:
                print("Warning: No footer table found, cannot set footer")
            
        except Exception as e:
            print(f"Warning: Could not setup footer: {e}")
            import traceback
            traceback.print_exc()
    
    def _apply_synergie_styling(self, doc: Document):
        """Apply Synergie styling to the entire document"""
        try:
            # Apply consistent styling to all paragraphs
            for para in doc.paragraphs:
                for run in para.runs:
                    # Always set font to Poppins (override any existing font)
                    run.font.name = self.font_name
                    
                    # Ensure proper color
                    if run.font.color.rgb is None:
                        run.font.color.rgb = self.black
            
            # Also apply to all sections (headers/footers)
            for section in doc.sections:
                # Header font
                for para in section.header.paragraphs:
                    for run in para.runs:
                        run.font.name = self.font_name
                # Footer font
                for para in section.footer.paragraphs:
                    for run in para.runs:
                        run.font.name = self.font_name
                # First page header/footer (if different)
                if section.different_first_page_header_footer:
                    for para in section.first_page_header.paragraphs:
                        for run in para.runs:
                            run.font.name = self.font_name
                    for para in section.first_page_footer.paragraphs:
                        for run in para.runs:
                            run.font.name = self.font_name
            
            print("Synergie styling applied (Poppins font enforced)")
            
        except Exception as e:
            print(f"Error applying styling: {e}")

if __name__ == "__main__":
    # Test the generator
    generator = TemplateResumeGeneratorV2("2.docx")
    
    # Test data
    test_cv_data = {
        'personal_info': {
            'full_name': 'Jeffrey C.F. van Aalst',
            'location': 'Purmerend',
            'birth_year': None,
            'phone': None,
            'email': None
        },
        'work_experience': [
            {
                'company': 'Zuidplus B.V.',
                'position': 'Coördinator Opleverproces',
                'start_date': '2018',
                'end_date': None,
                'is_current': True,
                'projects': [
                    {
                        'name': 'Zuidasdok',
                        'client': 'Zuidplus',
                        'period': '2018-heden',
                        'role': 'Coördinator Opleverproces',
                        'responsibilities': [
                            'Coördineren van het opleverproces',
                            'Beheren van documentatie',
                            'Communicatie met stakeholders'
                        ]
                    }
                ]
            }
        ],
        'education': [
            {
                'degree': 'HBO Bouwkunde',
                'institution': 'Hogeschool van Amsterdam',
                'period': '2005-2009'
            }
        ],
        'skills': ['Primavera P6', 'MS Project', 'Planning Management'],
        'profile_summary': 'Ervaren projectmanager met focus op infrastructuur en bouwprojecten.',
        'confidence_score': 0.9
    }
    
    print("Testing template-based Resumé generation V2...")
    result = generator.generate_resume(test_cv_data, "test_output_v2.docx")
    
    if result['success']:
        print("✅ Resumé generation successful")
        print(f"Output file: {result['output_path']}")
        print(f"File size: {result['file_size']} bytes")
    else:
        print(f"❌ Resumé generation failed: {result['error']}")
